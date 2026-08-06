from __future__ import annotations

from pathlib import Path

from nodeautomationtoolkit.core.definition import node
from nodeautomationtoolkit.core.word_types import (
    WordDocument,
    WordParagraph,
    WordParagraphs,
    WordSaveResult,
)


def _document_class():
    from docx import Document

    return Document


def _validated_docx_path(path: str, *, must_exist: bool) -> Path:
    if not path.strip():
        raise ValueError("Шлях до DOCX не вказано")
    resolved = Path(path).expanduser()
    if resolved.suffix.casefold() != ".docx":
        resolved = resolved.with_suffix(".docx")
    if must_exist and not resolved.is_file():
        raise FileNotFoundError(f"DOCX не знайдено: {resolved}")
    return resolved


def _iter_all_paragraphs(document):
    seen_cells: set[int] = set()

    def from_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_id = id(cell._tc)
                    if cell_id in seen_cells:
                        continue
                    seen_cells.add(cell_id)
                    yield from cell.paragraphs
                    yield from from_tables(cell.tables)

    yield from document.paragraphs
    yield from from_tables(document.tables)


@node(
    name="Прочитати DOCX",
    category="Word",
    description="Читає локальний DOCX, його назву, текст і абзаци.",
    type_id="builtin.word.read_docx",
    outputs={
        "document": "WordDocument",
        "file_name": "str",
        "paragraphs": "WordParagraphs",
        "text": "str",
    },
)
def read_docx(path: str) -> dict:
    source_path = _validated_docx_path(path, must_exist=True)
    document = _document_class()(source_path)
    raw_paragraphs = list(_iter_all_paragraphs(document))
    paragraph_items = tuple(
        WordParagraph(
            index=index,
            text=paragraph.text,
            style=paragraph.style.name if paragraph.style is not None else "",
            is_empty=not paragraph.text.strip(),
        )
        for index, paragraph in enumerate(raw_paragraphs)
    )
    paragraphs = WordParagraphs(paragraph_items, str(source_path))
    text = paragraphs.text()
    word_document = WordDocument(
        path=str(source_path),
        file_name=source_path.name,
        paragraphs=paragraphs,
        text=text,
    )
    return {
        "document": word_document,
        "file_name": word_document.file_name,
        "paragraphs": paragraphs,
        "text": text,
    }


@node(
    name="Абзаци документа",
    category="Word",
    description="Повертає структуровані абзаци прочитаного Word-документа.",
    type_id="builtin.word.document_paragraphs",
)
def document_paragraphs(document: WordDocument) -> WordParagraphs:
    return document.paragraphs


@node(
    name="Текст документа",
    category="Word",
    description="Повертає суцільний текст прочитаного Word-документа.",
    type_id="builtin.word.document_text",
)
def document_text(document: WordDocument, include_empty: bool = False) -> str:
    if include_empty:
        return document.text
    return "\n".join(item.text for item in document.paragraphs if not item.is_empty)


@node(
    name="Знайти абзаци",
    category="Word",
    description="Вибирає абзаци, у тексті яких є вказаний збіг.",
    type_id="builtin.word.filter_paragraphs",
)
def filter_paragraphs(
    paragraphs: WordParagraphs,
    contains: str,
    ignore_case: bool = True,
) -> WordParagraphs:
    if not contains:
        return paragraphs
    needle = contains.casefold() if ignore_case else contains
    selected = []
    for paragraph in paragraphs:
        candidate = paragraph.text.casefold() if ignore_case else paragraph.text
        if needle in candidate:
            selected.append(paragraph)
    return WordParagraphs(tuple(selected), paragraphs.source_path)


@node(
    name="Абзаци в текст",
    category="Word",
    description="Об'єднує вибрані Word-абзаци у текст.",
    type_id="builtin.word.paragraphs_to_text",
)
def paragraphs_to_text(paragraphs: WordParagraphs, separator: str = "\n") -> str:
    return paragraphs.text(separator)


@node(
    name="Створити DOCX",
    category="Word",
    description="Створює новий Word-документ із тексту й зберігає його локально.",
    type_id="builtin.word.create_docx",
    execution_inputs=("exec",),
    execution_outputs=("then",),
    preview_policy="never",
)
def create_docx(text: str, output_path: str, title: str = "") -> WordSaveResult:
    target_path = _validated_docx_path(output_path, must_exist=False)
    target_path.parent.mkdir(parents=True, exist_ok=True)
    document = _document_class()()
    if title.strip():
        document.add_heading(title.strip(), level=1)
    lines = text.splitlines() or [text]
    for line in lines:
        paragraph = document.add_paragraph(line)
        paragraph.paragraph_format.keep_together = True
    document.save(target_path)
    return WordSaveResult(
        path=str(target_path),
        paragraph_count=len(document.paragraphs),
    )


@node(
    name="Зберегти вибрані абзаци",
    category="Word",
    description=(
        "Створює копію вихідного DOCX, залишає вибрані абзаци та зберігає форматування."
    ),
    type_id="builtin.word.save_selected_paragraphs",
    execution_inputs=("exec",),
    execution_outputs=("then",),
    preview_policy="never",
)
def save_selected_paragraphs(
    document: WordDocument,
    paragraphs: WordParagraphs,
    output_path: str,
    keep_together: bool = True,
) -> WordSaveResult:
    target_path = _validated_docx_path(output_path, must_exist=False)
    target_path.parent.mkdir(parents=True, exist_ok=True)
    source = _document_class()(document.path)
    selected_indices = {item.index for item in paragraphs}

    for index, paragraph in reversed(list(enumerate(source.paragraphs))):
        if index not in selected_indices:
            element = paragraph._element
            element.getparent().remove(element)
        elif keep_together:
            paragraph.paragraph_format.keep_together = True

    for table in list(source.tables):
        element = table._element
        element.getparent().remove(element)

    source.save(target_path)
    return WordSaveResult(
        path=str(target_path),
        paragraph_count=len(selected_indices),
        message="Вибрані абзаци збережено",
    )


@node(
    name="Аналіз колонтитулів та нумерації",
    category="Word",
    description=(
        "Шукає верхні й нижні колонтитули, номери сторінок та службові поля "
        "у файлі DOCX. Виводить список колонтитулів, наявність нумерації та таблицю."
    ),
    type_id="builtin.word.analyze_headers_footers",
    outputs={
        "headers": "List",
        "footers": "List",
        "has_page_numbers": "bool",
        "details": "DataTable",
        "summary": "str",
    },
)
def analyze_headers_footers(path: str = "") -> dict:
    import re
    from nodeautomationtoolkit.core.table_types import DataTable

    source_path = _validated_docx_path(path, must_exist=True)
    doc = _document_class()(source_path)

    headers: list[str] = []
    footers: list[str] = []
    has_page_numbers = False
    table_rows = []

    page_regex = re.compile(
        r"(?i)\b(?:стор\.?|сторінка|page|numpages)\b|\b\d+\s*(?:з|/)\s*\d+\b"
    )

    for idx, section in enumerate(doc.sections, start=1):
        sec_headers = [
            ("Верхній (основний)", section.header),
            ("Верхній (перша ст.)", getattr(section, "first_page_header", None)),
            ("Верхній (парна ст.)", getattr(section, "even_page_header", None)),
        ]
        sec_footers = [
            ("Нижній (основний)", section.footer),
            ("Нижній (перша ст.)", getattr(section, "first_page_footer", None)),
            ("Нижній (парна ст.)", getattr(section, "even_page_footer", None)),
        ]

        for kind, h_obj in sec_headers:
            if h_obj is None or getattr(h_obj, "is_linked_to_previous", False):
                continue
            h_text = "\n".join(p.text for p in h_obj.paragraphs if p.text.strip()).strip()
            if h_text:
                headers.append(h_text)
            h_xml = h_obj._element.xml
            has_num = bool("PAGE" in h_xml or "NUMPAGES" in h_xml or page_regex.search(h_text))
            if has_num:
                has_page_numbers = True
            if h_text or has_num:
                table_rows.append((idx, kind, h_text or "(поля нумерації)", "Так" if has_num else "Ні"))

        for kind, f_obj in sec_footers:
            if f_obj is None or getattr(f_obj, "is_linked_to_previous", False):
                continue
            f_text = "\n".join(p.text for p in f_obj.paragraphs if p.text.strip()).strip()
            if f_text:
                footers.append(f_text)
            f_xml = f_obj._element.xml
            has_num = bool("PAGE" in f_xml or "NUMPAGES" in f_xml or page_regex.search(f_text))
            if has_num:
                has_page_numbers = True
            if f_text or has_num:
                table_rows.append((idx, kind, f_text or "(поля нумерації)", "Так" if has_num else "Ні"))

    headers = list(dict.fromkeys(headers))
    footers = list(dict.fromkeys(footers))
    table = DataTable(
        ("Секція", "Тип колонтитула", "Текст", "Нумерація сторінок"),
        tuple(table_rows),
        "Аналіз колонтитулів",
    )
    summary = (
        f"Верхніх колонтитулів: {len(headers)} · Нижніх: {len(footers)} · "
        f"Нумерація: {'Знайдено' if has_page_numbers else 'Відсутня'}"
    )
    return {
        "headers": headers,
        "footers": footers,
        "has_page_numbers": has_page_numbers,
        "details": table,
        "summary": summary,
    }


@node(
    name="Видалити колонтитули та нумерацію",
    category="Word",
    description=(
        "Видаляє всі верхні/нижні колонтитули та поля нумерації сторінок із файлу DOCX."
    ),
    type_id="builtin.word.clear_headers_footers",
    execution_inputs=("exec",),
    execution_outputs=("then",),
    outputs={
        "path": "str",
        "cleared_sections": "int",
        "summary": "str",
    },
)
def clear_headers_footers(
    path: str = "",
    output_path: str = "",
    clear_headers: bool = True,
    clear_footers: bool = True,
) -> dict:
    source_path = _validated_docx_path(path, must_exist=True)
    target_path = _validated_docx_path(output_path or str(source_path), must_exist=False)
    target_path.parent.mkdir(parents=True, exist_ok=True)
    doc = _document_class()(source_path)

    cleared_count = 0
    for section in doc.sections:
        cleared_count += 1
        if clear_headers:
            for h in [section.header, getattr(section, "first_page_header", None), getattr(section, "even_page_header", None)]:
                if h is not None:
                    for p in h.paragraphs:
                        p.text = ""
        if clear_footers:
            for f in [section.footer, getattr(section, "first_page_footer", None), getattr(section, "even_page_footer", None)]:
                if f is not None:
                    for p in f.paragraphs:
                        p.text = ""

    doc.save(target_path)
    summary = f"Очищено колонтитули у {cleared_count} секціях -> {target_path.name}"
    return {
        "path": str(target_path),
        "cleared_sections": cleared_count,
        "summary": summary,
    }


@node(
    name="Замінити текст у DOCX",
    category="Word",
    description="Знаходить та замінює всі входження тексту у DOCX документі, зберігаючи форматування.",
    type_id="builtin.word.replace_in_docx",
    execution_inputs=("exec",),
    execution_outputs=("then",),
    preview_policy="never",
    outputs={
        "path": "str",
        "replaced_count": "int",
        "summary": "str",
    },
)
def replace_in_docx(
    path: str = "",
    old_text: str = "",
    new_text: str = "",
    output_path: str = "",
    ignore_case: bool = False,
) -> dict:
    from docx import Document
    source_path = _validated_docx_path(path, must_exist=True)
    target_path = _validated_docx_path(output_path or str(source_path), must_exist=False)
    target_path.parent.mkdir(parents=True, exist_ok=True)
    
    doc = Document(source_path)
    replaced_count = 0
    
    for p in doc.paragraphs:
        for run in p.runs:
            if ignore_case:
                if old_text.lower() in run.text.lower():
                    # For simplicity, if ignore case, only replace if match, but simple replace might be tricky
                    # Let's just do simple replace with ignore_case logic
                    import re
                    matches = len(re.findall(re.escape(old_text), run.text, flags=re.IGNORECASE))
                    if matches > 0:
                        replaced_count += matches
                        run.text = re.sub(re.escape(old_text), new_text, run.text, flags=re.IGNORECASE)
            else:
                if old_text in run.text:
                    replaced_count += run.text.count(old_text)
                    run.text = run.text.replace(old_text, new_text)
                
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if ignore_case:
                            import re
                            matches = len(re.findall(re.escape(old_text), run.text, flags=re.IGNORECASE))
                            if matches > 0:
                                replaced_count += matches
                                run.text = re.sub(re.escape(old_text), new_text, run.text, flags=re.IGNORECASE)
                        else:
                            if old_text in run.text:
                                replaced_count += run.text.count(old_text)
                                run.text = run.text.replace(old_text, new_text)
                            
    doc.save(target_path)
    summary = f"Замінено {replaced_count} входжень -> {target_path.name}"
    return {
        "path": str(target_path),
        "replaced_count": replaced_count,
        "summary": summary,
    }


@node(
    name="Злити DOCX файли",
    category="Word",
    description="Об'єднує декілька файлів .docx в один.",
    type_id="builtin.word.merge_docx",
    execution_inputs=("exec",),
    execution_outputs=("then",),
    preview_policy="never",
    outputs={
        "path": "str",
        "merged_count": "int",
        "summary": "str",
    },
)
def merge_docx(
    files: list = None,
    output_path: str = "",
    add_page_break: bool = True,
) -> dict:
    from docx import Document
    if not files:
        return {"path": "", "merged_count": 0, "summary": "Немає файлів для об'єднання."}
        
    target_path = _validated_docx_path(output_path, must_exist=False)
    target_path.parent.mkdir(parents=True, exist_ok=True)
    
    merged_count = 0
    merged_doc = None
    
    for file_path in files:
        path_obj = _validated_docx_path(file_path, must_exist=True)
        if merged_doc is None:
            merged_doc = Document(path_obj)
            merged_count += 1
        else:
            if add_page_break:
                merged_doc.add_page_break()
            sub_doc = Document(path_obj)
            for p in sub_doc.paragraphs:
                # Copying just text for simplicity, python-docx doesn't easily support full element copying
                merged_doc.add_paragraph(p.text)
            merged_count += 1
            
    if merged_doc:
        merged_doc.save(target_path)
        
    summary = f"Об'єднано {merged_count} файлів -> {target_path.name}"
    return {
        "path": str(target_path),
        "merged_count": merged_count,
        "summary": summary,
    }


@node(
    name="Підрахунок слів у DOCX",
    category="Word",
    description="Підраховує кількість слів, символів, абзаців та сторінок (приблизно) у DOCX документі.",
    type_id="builtin.word.word_count",
    outputs={
        "words": "int",
        "characters": "int",
        "paragraphs": "int",
        "pages_estimate": "int",
        "summary": "str",
    },
)
def word_count(
    path: str = "",
) -> dict:
    from docx import Document
    source_path = _validated_docx_path(path, must_exist=True)
    doc = Document(source_path)
    
    words = 0
    characters = 0
    paragraphs = 0
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs += 1
            words += len(text.split())
            characters += len(text.replace(" ", ""))
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text.strip()
                    if text:
                        paragraphs += 1
                        words += len(text.split())
                        characters += len(text.replace(" ", ""))
                        
    pages_estimate = max(1, words // 300)
    
    summary = f"Слів: {words}, Символів: {characters}, Абзаців: {paragraphs}, Сторінок: ~{pages_estimate}"
    return {
        "words": words,
        "characters": characters,
        "paragraphs": paragraphs,
        "pages_estimate": pages_estimate,
        "summary": summary,
    }


@node(
    name="Метадані документа",
    category="Word",
    description="Зчитує або встановлює метадані (заголовок, автор, тема) DOCX документа.",
    type_id="builtin.word.set_metadata",
    execution_inputs=("exec",),
    execution_outputs=("then",),
    preview_policy="never",
    outputs={
        "path": "str",
        "summary": "str",
    },
)
def set_metadata(
    path: str = "",
    title: str = "",
    author: str = "",
    subject: str = "",
    output_path: str = "",
) -> dict:
    from docx import Document
    source_path = _validated_docx_path(path, must_exist=True)
    target_path = _validated_docx_path(output_path or str(source_path), must_exist=False)
    target_path.parent.mkdir(parents=True, exist_ok=True)
    
    doc = Document(source_path)
    core_properties = doc.core_properties
    
    if title:
        core_properties.title = title
    if author:
        core_properties.author = author
    if subject:
        core_properties.subject = subject
        
    doc.save(target_path)
    summary = f"Метадані оновлено -> {target_path.name}"
    return {
        "path": str(target_path),
        "summary": summary,
    }
