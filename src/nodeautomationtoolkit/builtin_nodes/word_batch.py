from __future__ import annotations

import json
import os
import re
import shutil
import tempfile
import uuid
from pathlib import Path

from nodeautomationtoolkit.core.batch_types import DocumentVariant, WordDocumentBatch
from nodeautomationtoolkit.core.definition import node


def _required_docx(path: str, label: str) -> Path:
    candidate = Path(path).expanduser()
    if not path.strip() or not candidate.is_file():
        raise FileNotFoundError(f"{label} не знайдено: {path or '(шлях порожній)'}")
    if candidate.suffix.casefold() != ".docx":
        raise ValueError(f"{label} має бути DOCX: {candidate}")
    return candidate.resolve()


def _parse_variants(
    groups: dict | None,
    names: list | None,
    names_text: str,
    fields_json: str,
) -> tuple[DocumentVariant, ...]:
    variants: list[DocumentVariant] = []
    if groups:
        for name, content in groups.items():
            if isinstance(content, dict):
                values = dict(content)
                values.setdefault("name", str(name))
                values.setdefault("marker", str(name))
                if str(values.get("content", "")).strip():
                    variants.append(DocumentVariant(str(values["name"]), tuple(values.items())))
            elif str(content).strip():
                variants.append(
                    DocumentVariant(str(name), (("marker", str(name)), ("content", content)))
                )
    elif fields_json.strip():
        payload = json.loads(fields_json)
        if isinstance(payload, dict):
            for name, fields in payload.items():
                values = fields if isinstance(fields, dict) else {"value": fields}
                variants.append(DocumentVariant(str(name), tuple(values.items())))
        elif isinstance(payload, list):
            for index, item in enumerate(payload, start=1):
                if isinstance(item, dict):
                    values = dict(item)
                    name = str(values.pop("name", f"Документ {index}"))
                    variants.append(DocumentVariant(name, tuple(values.items())))
                else:
                    variants.append(DocumentVariant(str(item)))
        else:
            raise ValueError("Поля варіантів мають бути JSON-об'єктом або списком")
    elif names:
        variants = [DocumentVariant(str(item)) for item in names if str(item).strip()]
    else:
        variants = [
            DocumentVariant(line.strip()) for line in names_text.splitlines() if line.strip()
        ]
    if not variants:
        variants = [DocumentVariant("Витяг")]
    return tuple(variants)


@node(
    name="Створити пакет документів",
    category="Word · Пакет",
    description=(
        "Бере один DOCX і створює план декількох вихідних документів. "
        "Назви можна подати списком, рядками або JSON із полями."
    ),
    type_id="builtin.word_batch.create",
)
def create_document_batch(
    source_path: str = "",
    template_path: str = "",
    groups: dict | None = None,
    names: list | None = None,
    names_text: str = "Витяг",
    fields_json: str = "",
) -> WordDocumentBatch:
    source = _required_docx(
        template_path or source_path,
        "DOCX-заготовка" if template_path else "Вхідний документ",
    )
    return WordDocumentBatch(
        source_path=str(source),
        variants=_parse_variants(groups, names, names_text, fields_json),
    )


@node(
    name="Замінити текст у пакеті",
    category="Word · Пакет",
    description=(
        "Додає заміну для кожного документа. У новому тексті працюють поля {{name}} та поля з JSON."
    ),
    type_id="builtin.word_batch.replace_text",
)
def batch_replace_text(
    batch: WordDocumentBatch,
    find: str = "",
    replacement_text: str = "",
    ignore_case: bool = True,
    replace_all: bool = True,
) -> WordDocumentBatch:
    if not find:
        raise ValueError("Не вказано текст для пошуку")
    return batch.with_operation(
        "replace_text",
        find=find,
        replacement_text=replacement_text,
        ignore_case=ignore_case,
        replace_all=replace_all,
    )


@node(
    name="Заповнити плейсхолдер",
    category="Word · Пакет",
    description=(
        "Шукає мітку на кшталт {{ТЕКСТ}} і вставляє заданий багаторядковий текст "
        "у кожен документ пакета."
    ),
    type_id="builtin.word_batch.fill_placeholder",
)
def batch_fill_placeholder(
    batch: WordDocumentBatch,
    placeholder: str = "{{ТЕКСТ}}",
    replacement_text: str = "",
) -> WordDocumentBatch:
    if not placeholder:
        raise ValueError("Плейсхолдер не може бути порожнім")
    return batch.with_operation(
        "replace_text",
        find=placeholder,
        replacement_text=replacement_text,
        ignore_case=False,
        replace_all=True,
    )


@node(
    name="Замінити основну частину",
    category="Word · Пакет",
    description=(
        "У кожній копії наказу знаходить початок параграфів/причин/пунктів та "
        "межу перед підписом, після чого вставляє згрупований {{content}}."
    ),
    type_id="builtin.word_batch.replace_body",
)
def batch_replace_body(
    batch: WordDocumentBatch,
    start_pattern: str = (
        r"(?im)^\s*(?:§\s*\d+|параграф\s+\d+|розділ\s+[IVXLC\d]+|"
        r"відповідно|підстава|причина|у\s+зв'язку|на\s+підставі|"
        r"на\s+виконання|звільнити|призначити|перемістити|зарахувати|"
        r"виключити|\d+(?:\.\d+)*[.)])"
    ),
    end_pattern: str = (
        r"(?im)^\s*(?:командир|командувач|начальник|заступник|керівник|голова|директор|"
        r"т\.в\.о\.|тимчасово\s+виконуючий)\b"
    ),
    replacement_text: str = "{{content}}",
) -> WordDocumentBatch:
    _compile_block_pattern(start_pattern, "початку")
    _compile_block_pattern(end_pattern, "кінця")
    return batch.with_operation(
        "replace_regex_block",
        start_pattern=start_pattern,
        end_pattern=end_pattern,
        replacement_text=replacement_text,
    )


@node(
    name="Остання сторінка з прикладу",
    category="Word · Пакет",
    description=(
        "Видаляє останню сторінку кожного документа й вставляє замість неї "
        "вміст указаного DOCX-прикладу. Потребує Microsoft Word."
    ),
    type_id="builtin.word_batch.replace_last_page",
)
def batch_replace_last_page(
    batch: WordDocumentBatch,
    example_page_path: str = "",
) -> WordDocumentBatch:
    example = _required_docx(example_page_path, "Приклад останньої сторінки")
    return batch.with_operation("replace_last_page", example_page_path=str(example))


def _normalize_page_selector(page: str | int) -> str | int:
    if isinstance(page, int):
        if page < 1:
            raise ValueError("Номер сторінки має бути більшим за нуль")
        return page
    value = str(page).strip().casefold()
    first_aliases = {"перша", "першу", "першої", "first", "початок", "на початку"}
    last_aliases = {"остання", "останню", "останньої", "last", "кінець", "в кінці"}
    if value in first_aliases:
        return "first"
    if value in last_aliases:
        return "last"
    try:
        page_number = int(value)
    except ValueError as error:
        raise ValueError("Сторінка має бути числом, словом 'перша' або словом 'остання'") from error
    if page_number < 1:
        raise ValueError("Номер сторінки має бути більшим за нуль")
    return page_number


@node(
    name="Видалити сторінку",
    category="Word · Пакет",
    description=(
        "Видаляє одну сторінку в кожному документі пакета. У полі page вкажіть "
        "номер, 'перша' або 'остання'. Точні межі сторінок визначає Microsoft Word."
    ),
    type_id="builtin.word_batch.delete_page",
)
def batch_delete_page(
    batch: WordDocumentBatch,
    page: str = "остання",
) -> WordDocumentBatch:
    return batch.with_operation("delete_page", page=_normalize_page_selector(page))


@node(
    name="Вставити сторінку",
    category="Word · Пакет",
    description=(
        "Вставляє нову сторінку в кожен документ пакета. 'перша' додає її на "
        "початок, 'остання' — у кінець, а номер — перед сторінкою з цим номером. "
        "Поле page_docx можна лишити порожнім для чистої сторінки або вказати "
        "DOCX-заготовку з готовим оформленням. Потребує Microsoft Word."
    ),
    type_id="builtin.word_batch.insert_page",
)
def batch_insert_page(
    batch: WordDocumentBatch,
    page: str = "остання",
    page_docx: str = "",
) -> WordDocumentBatch:
    template = ""
    if page_docx.strip():
        template = str(_required_docx(page_docx, "DOCX-заготовка сторінки"))
    return batch.with_operation(
        "insert_page",
        page=_normalize_page_selector(page),
        page_docx=template,
    )


@node(
    name="Замінити підпис або блок",
    category="Word · Пакет",
    description=(
        "Знаходить початкову мітку підпису, видаляє блок до кінцевої мітки "
        "або кінця документа та вставляє текст чи DOCX-приклад."
    ),
    type_id="builtin.word_batch.replace_signature",
)
def batch_replace_signature(
    batch: WordDocumentBatch,
    start_marker: str = "",
    end_marker: str = "",
    replacement_text: str = "",
    replacement_docx: str = "",
) -> WordDocumentBatch:
    if not start_marker:
        raise ValueError("Не вказано початкову мітку блока підпису")
    if not replacement_text and not replacement_docx:
        raise ValueError("Вкажіть новий текст підпису або DOCX-приклад")
    template = ""
    if replacement_docx:
        template = str(_required_docx(replacement_docx, "Приклад підпису"))
    return batch.with_operation(
        "replace_block",
        start_marker=start_marker,
        end_marker=end_marker,
        replacement_text=replacement_text,
        replacement_docx=template,
    )


@node(
    name="Форматувати сторінки",
    category="Word · Пакет",
    description=(
        "Змінює шрифт, розмір, вирівнювання та інтервали лише на вказаних "
        "сторінках. 0 в останній сторінці означає кінець документа."
    ),
    type_id="builtin.word_batch.format_pages",
)
def batch_format_pages(
    batch: WordDocumentBatch,
    first_page: int = 1,
    last_page: int = 0,
    font_name: str = "",
    font_size: float = 0.0,
    line_spacing: float = 1.0,
    space_before: float = 0.0,
    space_after: float = 0.0,
    alignment: str = "не змінювати",
    keep_together: bool = True,
) -> WordDocumentBatch:
    if first_page < 1 or last_page < 0:
        raise ValueError("Номери сторінок мають бути додатними")
    if last_page and last_page < first_page:
        raise ValueError("Остання сторінка не може бути перед першою")
    if line_spacing <= 0:
        raise ValueError("Міжрядковий інтервал має бути більшим за нуль")
    return batch.with_operation(
        "format_pages",
        first_page=first_page,
        last_page=last_page,
        font_name=font_name,
        font_size=font_size,
        line_spacing=line_spacing,
        space_before=space_before,
        space_after=space_after,
        alignment=alignment,
        keep_together=keep_together,
    )


@node(
    name="Не розривати пункти",
    category="Word · Пакет",
    description=(
        "Позначає кожен нумерований пункт як неподільний блок, щоб Word не переносив "
        "половину пункту на наступну сторінку."
    ),
    type_id="builtin.word_batch.keep_items_together",
)
def batch_keep_items_together(
    batch: WordDocumentBatch,
    item_pattern: str = r"^\s*\d+(?:\.\d+)*[.)]\s*",
    signature_pattern: str = (
        r"^\s*(?:командир|командувач|начальник|заступник|керівник|голова|директор|"
        r"т\.в\.о\.|тимчасово\s+виконуючий)\b"
    ),
) -> WordDocumentBatch:
    _compile_item_pattern(item_pattern)
    _compile_item_pattern(signature_pattern)
    return batch.with_operation(
        "keep_items_together",
        item_pattern=item_pattern,
        signature_pattern=signature_pattern,
    )


def _expand_fields(text: str, variant: DocumentVariant) -> str:
    result = text
    for key, value in variant.values().items():
        result = result.replace("{{" + str(key) + "}}", str(value))
    return result


def _safe_file_stem(value: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", value).strip(" .")
    return cleaned or "Документ"


def _compile_item_pattern(pattern: str) -> re.Pattern:
    try:
        return re.compile(pattern)
    except re.error as error:
        raise ValueError(f"Некоректне правило пунктів: {error}") from error


def _compile_block_pattern(pattern: str, label: str) -> re.Pattern:
    try:
        return re.compile(pattern)
    except re.error as error:
        raise ValueError(f"Некоректне правило {label} блока: {error}") from error


def _output_paths(
    batch: WordDocumentBatch,
    output_folder: Path,
    filename_template: str,
    overwrite: bool,
) -> list[Path]:
    result = []
    used: set[str] = set()
    for variant in batch.variants:
        filename = _expand_fields(filename_template, variant)
        path = Path(filename)
        stem = _safe_file_stem(path.stem)
        suffix = ".docx"
        candidate = output_folder / f"{stem}{suffix}"
        index = 2
        while candidate.name.casefold() in used:
            candidate = output_folder / f"{stem} ({index}){suffix}"
            index += 1
        used.add(candidate.name.casefold())
        if candidate.exists() and not overwrite:
            raise FileExistsError(f"Файл уже існує: {candidate}")
        result.append(candidate)
    return result


def _find_word_range(document, text: str, start: int = 0, end: int | None = None):
    search_range = document.Range(Start=start, End=end or document.Content.End - 1)
    finder = search_range.Find
    finder.ClearFormatting()
    finder.Text = text
    finder.Forward = True
    finder.Wrap = 0
    finder.Format = False
    return search_range if finder.Execute() else None


def _replace_word_text(
    document,
    find_text: str,
    replacement_text: str,
    *,
    ignore_case: bool,
    replace_all: bool,
) -> int:
    count = 0
    search_start = 0
    while search_start < document.Content.End - 1:
        search_range = document.Range(Start=search_start, End=document.Content.End - 1)
        finder = search_range.Find
        finder.ClearFormatting()
        finder.Text = find_text
        finder.MatchCase = not ignore_case
        finder.Forward = True
        finder.Wrap = 0
        finder.Format = False
        if not finder.Execute():
            break
        search_range.Text = replacement_text.replace("\n", "\r")
        count += 1
        search_start = search_range.End
        if not replace_all:
            break
    return count


def _page_range(document, first_page: int, last_page: int):
    page_count = document.ComputeStatistics(2)
    if first_page > page_count:
        raise ValueError(f"У документі лише {page_count} сторінок")
    resolved_last = page_count if last_page == 0 else min(last_page, page_count)
    start = document.GoTo(What=1, Which=1, Count=first_page).Start
    if resolved_last < page_count:
        end = document.GoTo(What=1, Which=1, Count=resolved_last + 1).Start
    else:
        end = document.Content.End - 1
    return document.Range(Start=start, End=end)


def _resolve_page_number(document, selector: str | int, *, allow_append: bool = False) -> int:
    page_count = int(document.ComputeStatistics(2))
    normalized = _normalize_page_selector(selector)
    if normalized == "first":
        return 1
    if normalized == "last":
        return page_count + 1 if allow_append else page_count
    maximum = page_count + 1 if allow_append else page_count
    if normalized > maximum:
        if allow_append:
            raise ValueError(
                f"У документі {page_count} сторінок; вставити можна не далі сторінки "
                f"{page_count + 1}"
            )
        raise ValueError(f"У документі лише {page_count} сторінок")
    return normalized


def _insert_page(document, selector: str | int, page_docx: str = "") -> None:
    document.Repaginate()
    page_number = _resolve_page_number(document, selector, allow_append=True)
    page_count = int(document.ComputeStatistics(2))
    if page_number == page_count + 1:
        insertion_position = document.Content.End - 1
        document.Range(
            Start=insertion_position,
            End=insertion_position,
        ).InsertBreak(7)
        if page_docx:
            insertion_position = document.Content.End - 1
            document.Range(
                Start=insertion_position,
                End=insertion_position,
            ).InsertFile(page_docx)
    else:
        insertion_position = document.GoTo(
            What=1,
            Which=1,
            Count=page_number,
        ).Start
        document.Range(
            Start=insertion_position,
            End=insertion_position,
        ).InsertBreak(7)
        if page_docx:
            document.Range(
                Start=insertion_position,
                End=insertion_position,
            ).InsertFile(page_docx)
    document.Repaginate()


def _apply_word_operation(application, document, operation, variant: DocumentVariant) -> None:
    options = operation.options()
    if operation.kind == "replace_text":
        replacement = _expand_fields(options["replacement_text"], variant)
        _replace_word_text(
            document,
            options["find"],
            replacement,
            ignore_case=options["ignore_case"],
            replace_all=options["replace_all"],
        )
        return
    if operation.kind == "replace_last_page":
        document.Repaginate()
        page_count = document.ComputeStatistics(2)
        target = _page_range(document, page_count, page_count)
        target.Delete()
        target.InsertFile(options["example_page_path"])
        return
    if operation.kind == "delete_page":
        document.Repaginate()
        page_number = _resolve_page_number(document, options["page"])
        _page_range(document, page_number, page_number).Delete()
        document.Repaginate()
        return
    if operation.kind == "insert_page":
        _insert_page(document, options["page"], options["page_docx"])
        return
    if operation.kind == "replace_block":
        start_range = _find_word_range(document, options["start_marker"])
        if start_range is None:
            raise ValueError(f"Не знайдено початок блока: {options['start_marker']}")
        start = start_range.Start
        if options["end_marker"]:
            end_range = _find_word_range(
                document,
                options["end_marker"],
                start=start_range.End,
            )
            if end_range is None:
                raise ValueError(f"Не знайдено кінець блока: {options['end_marker']}")
            end = end_range.End
        else:
            end = document.Content.End - 1
        target = document.Range(Start=start, End=end)
        target.Delete()
        if options["replacement_docx"]:
            target.InsertFile(options["replacement_docx"])
        else:
            target.Text = _expand_fields(options["replacement_text"], variant).replace("\n", "\r")
        return
    if operation.kind == "replace_regex_block":
        content = str(document.Content.Text)
        start_match = _compile_block_pattern(options["start_pattern"], "початку").search(content)
        if start_match is None:
            raise ValueError("Не знайдено початок основної частини документа")
        end_match = _compile_block_pattern(options["end_pattern"], "кінця").search(
            content,
            start_match.end(),
        )
        if end_match is None:
            raise ValueError("Не знайдено межу перед підписом")
        target = document.Range(Start=start_match.start(), End=end_match.start())
        replacement = _expand_fields(options["replacement_text"], variant)
        target.Text = replacement.replace("\n", "\r") + "\r"
        return
    if operation.kind == "format_pages":
        target = _page_range(document, options["first_page"], options["last_page"])
        if options["font_name"]:
            target.Font.Name = options["font_name"]
        if options["font_size"] > 0:
            target.Font.Size = options["font_size"]
        paragraph_format = target.ParagraphFormat
        spacing = options["line_spacing"]
        if abs(spacing - 1.0) < 0.001:
            paragraph_format.LineSpacingRule = 0
        elif abs(spacing - 1.5) < 0.001:
            paragraph_format.LineSpacingRule = 1
        elif abs(spacing - 2.0) < 0.001:
            paragraph_format.LineSpacingRule = 2
        else:
            paragraph_format.LineSpacingRule = 5
            paragraph_format.LineSpacing = application.LinesToPoints(spacing)
        paragraph_format.SpaceBefore = options["space_before"]
        paragraph_format.SpaceAfter = options["space_after"]
        paragraph_format.KeepTogether = bool(options["keep_together"])
        alignment = options["alignment"].strip().casefold()
        alignments = {
            "ліворуч": 0,
            "по центру": 1,
            "центр": 1,
            "праворуч": 2,
            "по ширині": 3,
        }
        if alignment in alignments:
            paragraph_format.Alignment = alignments[alignment]
        return
    if operation.kind == "keep_items_together":
        pattern = _compile_item_pattern(options["item_pattern"])
        signature = _compile_item_pattern(options["signature_pattern"])
        paragraphs = list(document.Paragraphs)
        starts = [
            index
            for index, paragraph in enumerate(paragraphs)
            if pattern.search(str(paragraph.Range.Text).rstrip("\r\x07"))
        ]
        for position, start in enumerate(starts):
            stop = starts[position + 1] if position + 1 < len(starts) else len(paragraphs)
            for index in range(start + 1, stop):
                text = str(paragraphs[index].Range.Text).rstrip("\r\x07")
                if signature.search(text):
                    stop = index
                    break
            for index in range(start, stop):
                paragraph_format = paragraphs[index].Range.ParagraphFormat
                paragraph_format.KeepTogether = True
                paragraph_format.KeepWithNext = index < stop - 1
        return
    raise ValueError(f"Невідома пакетна операція: {operation.kind}")


def _word_application():
    try:
        import win32com.client
    except ImportError as error:
        raise RuntimeError(
            "Для сторінок і точного форматування потрібні Windows, Microsoft Word та pywin32"
        ) from error
    try:
        application = win32com.client.DispatchEx("Word.Application")
    except Exception as error:  # noqa: BLE001 - COM activation boundary
        raise RuntimeError(
            "Не вдалося запустити Microsoft Word. Перевірте, чи він встановлений."
        ) from error
    application.Visible = False
    application.DisplayAlerts = 0
    return application


def _save_batch_with_word(
    batch: WordDocumentBatch,
    targets: list[Path],
) -> list[int]:
    source = Path(batch.source_path)
    temporary_paths = [
        target.parent / f".nat-{uuid.uuid4().hex}-{target.name}" for target in targets
    ]
    application = _word_application()
    paragraph_counts: list[int] = []
    try:
        for temporary, variant in zip(temporary_paths, batch.variants, strict=True):
            shutil.copy2(source, temporary)
            document = application.Documents.Open(
                str(temporary),
                ReadOnly=False,
                AddToRecentFiles=False,
            )
            try:
                for operation in batch.operations:
                    _apply_word_operation(application, document, operation, variant)
                paragraph_counts.append(int(document.Paragraphs.Count))
                document.Save()
            finally:
                document.Close(SaveChanges=False)
        for temporary, target in zip(temporary_paths, targets, strict=True):
            os.replace(temporary, target)
    finally:
        application.Quit()
        for temporary in temporary_paths:
            temporary.unlink(missing_ok=True)
    return paragraph_counts


def _replace_paragraph_text(
    paragraph,
    find_text: str,
    replacement_text: str,
    *,
    ignore_case: bool,
    replace_all: bool,
) -> int:
    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text or not find_text:
        return 0
    flags = re.IGNORECASE if ignore_case else 0
    matches = list(re.finditer(re.escape(find_text), full_text, flags))
    if not replace_all:
        matches = matches[:1]
    count = 0
    for match in reversed(matches):
        spans = []
        position = 0
        for index, run in enumerate(paragraph.runs):
            end = position + len(run.text)
            spans.append((index, position, end))
            position = end
        touched = [item for item in spans if item[1] < match.end() and item[2] > match.start()]
        if not touched:
            continue
        first_index, first_start, _first_end = touched[0]
        last_index, last_start, _last_end = touched[-1]
        first_run = paragraph.runs[first_index]
        last_run = paragraph.runs[last_index]
        prefix = first_run.text[: match.start() - first_start]
        suffix = last_run.text[match.end() - last_start :]
        if first_index == last_index:
            first_run.text = prefix + replacement_text + suffix
        else:
            first_run.text = prefix + replacement_text
            for index in range(first_index + 1, last_index):
                paragraph.runs[index].text = ""
            last_run.text = suffix
        count += 1
    return count


def _iter_document_paragraphs(document):
    seen_cells: set[int] = set()

    def table_paragraphs(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_id = id(cell._tc)
                    if cell_id in seen_cells:
                        continue
                    seen_cells.add(cell_id)
                    yield from cell.paragraphs
                    yield from table_paragraphs(cell.tables)

    yield from document.paragraphs
    yield from table_paragraphs(document.tables)
    for section in document.sections:
        for part in (section.header, section.footer):
            yield from part.paragraphs
            yield from table_paragraphs(part.tables)


def _save_batch_portable(batch: WordDocumentBatch, targets: list[Path]) -> list[int]:
    from docx import Document

    source_bytes = Path(batch.source_path).read_bytes()
    temporary_paths = [
        target.parent / f".nat-{uuid.uuid4().hex}-{target.name}" for target in targets
    ]
    paragraph_counts = []
    try:
        for temporary, variant in zip(temporary_paths, batch.variants, strict=True):
            temporary.write_bytes(source_bytes)
            document = Document(temporary)
            for operation in batch.operations:
                if operation.kind == "keep_items_together":
                    pattern = _compile_item_pattern(operation.options()["item_pattern"])
                    signature = _compile_item_pattern(operation.options()["signature_pattern"])
                    paragraphs = list(_iter_document_paragraphs(document))
                    starts = [
                        index
                        for index, paragraph in enumerate(paragraphs)
                        if pattern.search(paragraph.text)
                    ]
                    for position, start in enumerate(starts):
                        stop = (
                            starts[position + 1] if position + 1 < len(starts) else len(paragraphs)
                        )
                        for index in range(start + 1, stop):
                            if signature.search(paragraphs[index].text):
                                stop = index
                                break
                        for index in range(start, stop):
                            paragraphs[index].paragraph_format.keep_together = True
                            paragraphs[index].paragraph_format.keep_with_next = index < stop - 1
                    continue
                if operation.kind != "replace_text":
                    raise RuntimeError("Ця операція потребує встановленого Microsoft Word")
                options = operation.options()
                replacement = _expand_fields(options["replacement_text"], variant)
                remaining_one = not options["replace_all"]
                for paragraph in _iter_document_paragraphs(document):
                    replaced = _replace_paragraph_text(
                        paragraph,
                        options["find"],
                        replacement,
                        ignore_case=options["ignore_case"],
                        replace_all=not remaining_one,
                    )
                    if remaining_one and replaced:
                        break
            paragraph_counts.append(len(document.paragraphs))
            document.save(temporary)
        for temporary, target in zip(temporary_paths, targets, strict=True):
            os.replace(temporary, target)
    finally:
        for temporary in temporary_paths:
            temporary.unlink(missing_ok=True)
    return paragraph_counts


@node(
    name="Повне прев'ю у Word",
    category="Word · Пакет",
    description=(
        "Збирає перший документ пакета у тимчасовий DOCX і відкриває його в "
        "Microsoft Word з реальними полями, таблицями, відступами й сторінками."
    ),
    type_id="builtin.word_batch.preview_in_word",
    outputs={"path": "str", "name": "str", "message": "str"},
    execution_inputs=("exec",),
    preview_policy="never",
)
def preview_document_batch(batch: WordDocumentBatch) -> dict:
    if not batch.variants:
        raise ValueError("У пакеті немає документів для перегляду")
    preview_folder = Path(tempfile.gettempdir()) / "NodeAutomationToolkit" / "Preview"
    preview_folder.mkdir(parents=True, exist_ok=True)
    variant = batch.variants[0]
    preview_batch = WordDocumentBatch(
        source_path=batch.source_path,
        variants=(variant,),
        operations=batch.operations,
    )
    target = preview_folder / f"{_safe_file_stem(variant.name)}-{uuid.uuid4().hex[:8]}.docx"
    _save_batch_with_word(preview_batch, [target])
    if not hasattr(os, "startfile"):
        raise RuntimeError("Повне прев'ю потребує Windows і Microsoft Word")
    os.startfile(str(target))  # type: ignore[attr-defined]  # noqa: S606
    return {
        "path": str(target),
        "name": target.name,
        "message": "Прев'ю відкрито у Microsoft Word",
    }


@node(
    name="Зберегти пакет DOCX",
    category="Word · Пакет",
    description=(
        "Фінальна нода: за один запуск застосовує всі правила й створює всі DOCX у вибраній папці."
    ),
    type_id="builtin.word_batch.save",
    outputs={"paths": "List", "count": "int", "message": "str"},
    execution_inputs=("exec",),
    execution_outputs=("then",),
    preview_policy="never",
)
@node(
    name="Очистити колонтитули у пакеті",
    category="Word · Пакет",
    description=(
        "Видаляє всі верхні/нижні колонтитули та поля нумерації з кожного "
        "документа в пакеті."
    ),
    type_id="builtin.word_batch.clear_headers_footers",
)
def batch_clear_headers_footers(
    batch: WordDocumentBatch,
    clear_headers: bool = True,
    clear_footers: bool = True,
) -> WordDocumentBatch:
    return batch.with_operation(
        "clear_headers_footers",
        clear_headers=clear_headers,
        clear_footers=clear_footers,
    )


def save_document_batch(
    batch: WordDocumentBatch,
    output_folder: str = "",
    filename_template: str = "{{name}}.docx",
    overwrite: bool = False,
) -> dict:
    if not output_folder.strip():
        raise ValueError("Не вибрано папку для збереження")
    destination = Path(output_folder).expanduser().resolve()
    destination.mkdir(parents=True, exist_ok=True)
    targets = _output_paths(batch, destination, filename_template, overwrite)
    portable_operations = {"replace_text", "keep_items_together"}
    requires_word = any(operation.kind not in portable_operations for operation in batch.operations)
    requires_word = requires_word or any(
        operation.kind == "replace_text"
        and any(
            "\n"
            in _expand_fields(
                operation.options()["replacement_text"],
                variant,
            )
            for variant in batch.variants
        )
        for operation in batch.operations
    )
    if requires_word:
        _save_batch_with_word(batch, targets)
    else:
        _save_batch_portable(batch, targets)
    return {
        "paths": [str(path) for path in targets],
        "count": len(targets),
        "message": f"Створено документів: {len(targets)}",
    }


@node(
    name="Сортувати накази по папках",
    category="Word · Пакет",
    description=(
        "Сканує папку з наказами DOCX, розпізнає номер і дату кожного наказу, "
        "створює підпапки за номерами наказів, переміщує файли та перейменовує "
        "за шаблоном 'прим_2_[Дата]_[Номер]'."
    ),
    type_id="builtin.word_batch.organize_by_number",
    outputs={
        "summary": "str",
        "processed_count": "int",
        "details": "DataTable",
    },
    execution_inputs=("exec",),
    execution_outputs=("then",),
    preview_policy="never",
)
def organize_orders_by_number(
    input_folder: str = "",
    create_subfolders: bool = True,
    rename_pattern: str = "прим_2_{date}_№{number}",
    copy_number: str = "2",
) -> dict:
    import shutil
    import docx
    from nodeautomationtoolkit.builtin_nodes.text_analysis import extract_order_fields
    from nodeautomationtoolkit.core.table_types import DataTable

    folder = Path(input_folder).expanduser()
    if not folder.is_dir():
        raise NotADirectoryError(f"Папку не знайдено: {input_folder}")

    files = [f for f in folder.glob("*.docx") if not f.name.startswith("~$")]
    rows = []
    processed = 0

    for doc_path in files:
        try:
            doc = docx.Document(doc_path)
            text = "\n".join(p.text for p in doc.paragraphs)
            extracted = extract_order_fields(text)
            order_num = extracted.get("order_number") or "б-н"
            order_date = extracted.get("order_date") or "б-д"

            clean_num = order_num.replace("/", "-").replace("\\", "-").strip()
            target_dir = folder / f"Наказ № {clean_num}" if create_subfolders else folder
            target_dir.mkdir(parents=True, exist_ok=True)

            formatted_name = rename_pattern.format(
                date=order_date.replace(" ", "_"),
                number=clean_num,
                copy=copy_number,
            ) + ".docx"

            new_file_path = target_dir / formatted_name
            if new_file_path.exists() and new_file_path != doc_path:
                new_file_path.unlink()

            moved_path = shutil.move(str(doc_path), str(new_file_path))
            processed += 1
            rows.append([doc_path.name, clean_num, order_date, Path(moved_path).name, str(target_dir.name)])
        except Exception as err:
            rows.append([doc_path.name, "Помилка", "", str(err), ""])

    table = DataTable(
        columns=["Початковий файл", "Номер наказу", "Дата наказу", "Нова назва", "Папка"],
        rows=rows,
    )
    return {
        "summary": f"Опрацьовано наказів: {processed} з {len(files)}",
        "processed_count": processed,
        "details": table,
    }
def _try_create_extracts_via_word_com(
    source_order_path: str,
    units_data: dict,
    out_dir: Path,
    template_docx_path: str,
    order_number: str,
    order_date: str,
    signatory_title: str,
    signatory_rank_name: str,
    certify_extract: bool,
    copy_number: str,
    executor_info: str,
    commander_signature_text: str,
    save_individual_files: bool,
) -> dict | None:
    """Гібридна обробка через MS Word COM Automation (нативне копіювання з 100% збереженням відступів, списків і форматування)."""
    try:
        import win32com.client
    except ImportError:
        return None

    if not source_order_path or not Path(source_order_path).is_file():
        return None

    word = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
    except Exception:
        return None

    try:
        abs_src = str(Path(source_order_path).resolve())
        src_doc = word.Documents.Open(abs_src)

        com_para_map = {}
        for i in range(1, src_doc.Paragraphs.Count + 1):
            p = src_doc.Paragraphs(i)
            t = p.Range.Text.strip()
            if t and t not in com_para_map:
                com_para_map[t] = p

        created_paths = []
        table_rows = []

        for unit_code, unit_info in units_data.items():
            clean_unit = str(unit_code).replace("/", "-").replace("\\", "-").strip()
            out_name = f"Витяг_{clean_unit}_№{order_number}.docx"
            out_file = out_dir / out_name
            abs_out = str(out_file.resolve())

            recipient_to = unit_info.get("recipient_to", "")
            destination_where = unit_info.get("destination_where", "")
            header_lines = unit_info.get("header_lines", [])
            items = unit_info.get("items", [])

            target_doc = word.Documents.Add()
            target_doc.PageSetup.TopMargin = 56.7   # 2.0 cm
            target_doc.PageSetup.BottomMargin = 56.7 # 2.0 cm
            target_doc.PageSetup.LeftMargin = 56.7   # 2.0 cm
            target_doc.PageSetup.RightMargin = 42.55 # 1.5 cm

            if recipient_to:
                p = target_doc.Paragraphs.Add()
                p.Range.Text = recipient_to + "\n"
                p.Range.ParagraphFormat.Alignment = 2  # wdAlignParagraphRight
                p.Range.Font.Bold = 1
                p.Range.Font.Size = 14

            if destination_where and destination_where.upper() not in ("КУДИ", "[КУДИ]"):
                p = target_doc.Paragraphs.Add()
                p.Range.Text = destination_where + "\n"
                p.Range.ParagraphFormat.Alignment = 2
                p.Range.Font.Italic = 1
                p.Range.Font.Size = 12
            else:
                p = target_doc.Paragraphs.Add()
                p.Range.Text = "КУДИ\n"
                p.Range.ParagraphFormat.Alignment = 2
                p.Range.Font.Bold = 1
                p.Range.Font.Size = 12
                try:
                    p.Range.Font.Color = 255  # wdColorRed
                except Exception:
                    pass

            if header_lines:
                for idx, line in enumerate(header_lines):
                    p = target_doc.Paragraphs.Add()
                    p.Range.Text = line + "\n"
                    p.Range.ParagraphFormat.Alignment = 1 if idx < len(header_lines) - 1 else 0
                    p.Range.Font.Bold = 1 if idx == 0 else 0
                    p.Range.Font.Size = 14
            else:
                p1 = target_doc.Paragraphs.Add()
                p1.Range.Text = "НАКАЗ КОМАНДИРА ВІЙСЬКОВОЇ ЧАСТИНИ А0000\n"
                p1.Range.ParagraphFormat.Alignment = 1
                p1.Range.Font.Bold = 1
                p1.Range.Font.Size = 14

                p2 = target_doc.Paragraphs.Add()
                p2.Range.Text = f"{order_date}               м. Львів               № {order_number}\n"
                p2.Range.ParagraphFormat.Alignment = 0
                p2.Range.Font.Size = 14

            p_title = target_doc.Paragraphs.Add()
            p_title.Range.Text = "ВИТЯГ З НАКАЗУ\n"
            p_title.Range.ParagraphFormat.Alignment = 1
            p_title.Range.Font.Bold = 1
            p_title.Range.Font.Size = 16

            if copy_number:
                p_copy = target_doc.Paragraphs.Add()
                p_copy.Range.Text = f"({copy_number})\n"
                p_copy.Range.ParagraphFormat.Alignment = 1
                p_copy.Range.Font.Italic = 1
                p_copy.Range.Font.Size = 12

            printed_headings = set()
            for item_data in items:
                heading = item_data.get("parent_heading", "") if isinstance(item_data, dict) else ""
                text = item_data.get("text", "") if isinstance(item_data, dict) else str(item_data)

                if heading and heading not in printed_headings:
                    src_h = com_para_map.get(heading.strip())
                    if src_h:
                        src_h.Range.Copy()
                        target_doc.Paragraphs(target_doc.Paragraphs.Count).Range.Paste()
                    else:
                        p_h = target_doc.Paragraphs.Add()
                        p_h.Range.Text = heading + "\n"
                        p_h.Range.Font.Bold = 1
                        p_h.Range.Font.Size = 14
                    printed_headings.add(heading)

                lines = [line.rstrip() for line in text.splitlines() if line.strip()]
                for line in lines:
                    src_p = com_para_map.get(line.strip())
                    if src_p:
                        src_p.Range.Copy()
                        target_doc.Paragraphs(target_doc.Paragraphs.Count).Range.Paste()
                    else:
                        p_line = target_doc.Paragraphs.Add()
                        p_line.Range.Text = line + "\n"
                        p_line.Range.Font.Size = 14

            if commander_signature_text:
                # Правило 5.3: рівно 2 порожні абзаци перед підписантом
                for _ in range(2):
                    p_blank = target_doc.Paragraphs.Add()
                    p_blank.Range.Text = "\n"
                    p_blank.Range.Font.Size = 14
                for sig_line in commander_signature_text.splitlines():
                    p_sig = target_doc.Paragraphs.Add()
                    p_sig.Range.Text = sig_line + "\n"
                    p_sig.Range.ParagraphFormat.Alignment = 2
                    p_sig.Range.Font.Bold = 1
                    p_sig.Range.Font.Size = 14

            if certify_extract:
                p_c1 = target_doc.Paragraphs.Add()
                p_c1.Range.Text = "Згідно з оригіналом:\n"
                p_c1.Range.Font.Bold = 1
                p_c1.Range.Font.Italic = 1
                p_c1.Range.Font.Size = 14

                p_c2 = target_doc.Paragraphs.Add()
                p_c2.Range.Text = signatory_title + "\n"
                p_c2.Range.Font.Size = 14

                p_c3 = target_doc.Paragraphs.Add()
                p_c3.Range.Text = signatory_rank_name + "\n"
                p_c3.Range.ParagraphFormat.Alignment = 2
                p_c3.Range.Font.Bold = 1
                p_c3.Range.Font.Size = 14

                year_str = order_date[-4:] if len(order_date) >= 4 and order_date[-4:].isdigit() else "2026"
                p_c4 = target_doc.Paragraphs.Add()
                p_c4.Range.Text = f"«____» ____________ {year_str} року\n"
                p_c4.Range.Font.Italic = 1
                p_c4.Range.Font.Size = 12

            if executor_info:
                p_exec = target_doc.Paragraphs.Add()
                p_exec.Range.Text = executor_info + "\n"
                p_exec.Range.Font.Italic = 0
                p_exec.Range.Font.Size = 8

            target_doc.SaveAs2(abs_out, FileFormat=16)
            target_doc.Close(SaveChanges=False)

            if save_individual_files:
                created_paths.append(str(out_file))

            table_rows.append((unit_code, len(items), out_name, str(out_file)))

        src_doc.Close(SaveChanges=False)

        clean_order_num = str(order_number).replace("/", "-").replace("\\", "-").strip()
        combined_path = out_dir / f"Всі_витяги_наказ_№{clean_order_num}.docx"
        final_paths = created_paths if save_individual_files else [str(combined_path)]

        table = DataTable(
            ("Військова частина", "Кількість пунктів", "Файл витягу", "Повний шлях"),
            tuple(table_rows),
        )

        return {
            "summary": f"Згенеровано витяги через MS Word COM Automation (нативне копіювання 100% форматування, {len(units_data)} адресатів)",
            "count": len(final_paths),
            "details": table,
            "paths": final_paths,
            "combined_path": str(combined_path),
        }
    except Exception:
        return None
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


@node(
    name="Пакетне створення витягів за ВЧ",
    category="Word · Пакет",
    description=(
        "Генерує окремі та 1 ЄДИНИЙ підсумковий DOCX з витягами з наказу для двостороннього друку (2 сторінки на 1 аркуш). "
        "Підтримує шаблон DOCX з плейсхолдерами {{кому}}, {{куди}}, {{зміст}}, {{номер_наказу}}, {{дата_наказу}}, {{засвідчення}}."
    ),
    type_id="builtin.order_batch.create_unit_extracts",
    outputs={
        "summary": "str",
        "count": "int",
        "details": "DataTable",
        "paths": "List",
        "combined_path": "str",
    },
    execution_inputs=("exec",),
    execution_outputs=("then",),
    preview_policy="never",
)
def create_unit_extracts(
    source_order_path: str = "",
    unit_paragraphs: dict | None = None,
    output_folder: str = "",
    template_docx_path: str = "",
    order_number: str = "",
    order_date: str = "",
    signatory_title: str = "Т.в.о. начальника штабу військової частини А0000",
    signatory_rank_name: str = "майор Петро СИДОРЕНКО",
    certify_extract: bool = True,
    copy_number: str = "прим. 2",
    executor_info: str = "Пупков ПУпенко 55-358",
    save_individual_files: bool = False,
    use_word_com: bool = True,
) -> dict:
    import copy
    import re
    import docx
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
    from nodeautomationtoolkit.builtin_nodes.text_analysis import extract_order_fields
    from nodeautomationtoolkit.core.table_types import DataTable

    out_dir = Path(output_folder).expanduser() if output_folder.strip() else Path("output/extracts")
    out_dir.mkdir(parents=True, exist_ok=True)

    commander_signature_text = ""
    if source_order_path.strip() and Path(source_order_path).is_file():
        src_file = Path(source_order_path)
        src_doc = docx.Document(source_order_path)
        full_text = "\n".join(p.text for p in src_doc.paragraphs)

        if not order_number or not order_date:
            extracted = extract_order_fields(full_text)
            order_number = order_number or extracted.get("order_number") or ""
            order_date = order_date or extracted.get("order_date") or ""

        file_stem = src_file.stem
        if not order_number:
            num_m = re.search(r"№\s*([A-Za-zА-Яа-яІіЇїЄєҐґ0-9./_-]+)", file_stem)
            if num_m:
                order_number = num_m.group(1).strip()
        if not order_date:
            date_m = re.search(r"(\d{1,2}\s+(?:січня|лютого|березня|квітня|травня|червня|липня|серпня|вересня|жовтня|листопада|грудня)\s+\d{4}(?:\s*року|\s*р\.)?)", file_stem, re.I)
            if not date_m:
                date_m = re.search(r"(\d{1,2}[./_]\d{1,2}[./_]\d{2,4})", file_stem)
            if date_m:
                order_date = date_m.group(1).strip().replace("_", ".")

        cmd_lines = []
        for p in reversed(src_doc.paragraphs):
            t = p.text.strip()
            if t and (
                "командир" in t.lower()
                or "начальник" in t.lower()
                or "командувач" in t.lower()
                or "генерал" in t.lower()
                or "полковник" in t.lower()
            ):
                cmd_lines.append(t)
                if len(cmd_lines) >= 2:
                    break
        if cmd_lines:
            commander_signature_text = "\n".join(reversed(cmd_lines))

    if use_word_com and source_order_path and Path(source_order_path).is_file():
        com_res = _try_create_extracts_via_word_com(
            source_order_path=source_order_path,
            units_data=unit_paragraphs or {},
            out_dir=out_dir,
            template_docx_path=template_docx_path,
            order_number=order_number or "б-н",
            order_date=order_date or "б-д",
            signatory_title=signatory_title,
            signatory_rank_name=signatory_rank_name,
            certify_extract=certify_extract,
            copy_number=copy_number,
            executor_info=executor_info,
            commander_signature_text=commander_signature_text,
            save_individual_files=save_individual_files,
        )
        if com_res is not None:
            return com_res

    # Побудова індексу: текст параграфа (без крайніх пробілів) → оригінальний об'єкт параграфа DOCX
    source_para_map: dict = {}
    if source_order_path and Path(source_order_path).is_file():
        try:
            _src_doc_map = docx.Document(source_order_path)
            for _sp in _src_doc_map.paragraphs:
                _k = _sp.text.strip()
                if _k and _k not in source_para_map:
                    source_para_map[_k] = _sp
        except Exception:
            pass

    order_number = order_number or "б-н"
    order_date = order_date or "б-д"

    units_data = unit_paragraphs or {}
    created_paths = []
    table_rows = []

    def _set_run_font(run, font_name="Times New Roman", size_pt=14, bold=False, italic=False, color_rgb=None):
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.italic = italic
        if color_rgb:
            try:
                run.font.color.rgb = docx.shared.RGBColor(*color_rgb)
            except Exception:
                pass

    def _add_styled_p(
        doc,
        text="",
        font_name="Times New Roman",
        size_pt=14,
        bold=False,
        italic=False,
        color_rgb=None,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=4,
        line_spacing=1.0,
        first_indent=1.25,
        left_indent=0,
        keep_with_next=False,
        keep_together=False,
    ):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = line_spacing
        if first_indent > 0:
            p.paragraph_format.first_line_indent = Cm(first_indent)
        if left_indent > 0:
            p.paragraph_format.left_indent = Cm(left_indent)
        if keep_with_next:
            p.paragraph_format.keep_with_next = True
        if keep_together:
            p.paragraph_format.keep_together = True
        if text:
            run = p.add_run(text)
            _set_run_font(run, font_name=font_name, size_pt=size_pt, bold=bold, italic=italic, color_rgb=color_rgb)
        return p

    def _insert_styled_p_before(
        target_p,
        text="",
        font_name="Times New Roman",
        size_pt=14,
        bold=False,
        italic=False,
        color_rgb=None,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=4,
        line_spacing=1.0,
        first_indent=1.25,
        left_indent=0,
        keep_with_next=False,
        keep_together=False,
    ):
        p = target_p.insert_paragraph_before()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = line_spacing
        if first_indent > 0:
            p.paragraph_format.first_line_indent = Cm(first_indent)
        if left_indent > 0:
            p.paragraph_format.left_indent = Cm(left_indent)
        if keep_with_next:
            p.paragraph_format.keep_with_next = True
        if keep_together:
            p.paragraph_format.keep_together = True
        if text:
            run = p.add_run(text)
            _set_run_font(run, font_name=font_name, size_pt=size_pt, bold=bold, italic=italic, color_rgb=color_rgb)
        return p

    def _apply_source_formatting(p, line_text, src_p=None, keep_with_next=False):
        """Зберігає форматування (відступи, рівнення, інтервали, рани) з оригіналу без зламу XML."""
        p.paragraph_format.keep_together = True
        if keep_with_next:
            p.paragraph_format.keep_with_next = True

        if src_p is not None:
            pf_src = src_p.paragraph_format
            pf_dst = p.paragraph_format

            if src_p.alignment is not None:
                p.alignment = src_p.alignment
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            if pf_src.first_line_indent is not None:
                pf_dst.first_line_indent = pf_src.first_line_indent
            elif not line_text.startswith("\t"):
                pf_dst.first_line_indent = Cm(1.25)

            if pf_src.left_indent is not None:
                pf_dst.left_indent = pf_src.left_indent
            if pf_src.right_indent is not None:
                pf_dst.right_indent = pf_src.right_indent

            if pf_src.space_after is not None:
                pf_dst.space_after = pf_src.space_after
            else:
                pf_dst.space_after = Pt(4)
            if pf_src.space_before is not None:
                pf_dst.space_before = pf_src.space_before
            if pf_src.line_spacing is not None:
                pf_dst.line_spacing = pf_src.line_spacing

            src_runs_text = "".join(r.text for r in src_p.runs).strip()
            if src_p.runs and src_runs_text == line_text.strip():
                for r in src_p.runs:
                    new_r = p.add_run(r.text)
                    new_r.font.name = r.font.name or "Times New Roman"
                    if r.font.size:
                        new_r.font.size = r.font.size
                    else:
                        new_r.font.size = Pt(14)
                    if r.bold is not None:
                        new_r.bold = r.bold
                    if r.italic is not None:
                        new_r.italic = r.italic
            else:
                r = p.add_run(line_text)
                _set_run_font(r, font_name="Times New Roman", size_pt=14, bold=False, italic=False)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.0
            if not line_text.startswith("\t"):
                p.paragraph_format.first_line_indent = Cm(1.25)
            r = p.add_run(line_text)
            _set_run_font(r, font_name="Times New Roman", size_pt=14, bold=False, italic=False)
        return p

    def _insert_para_from_source(target_p, line_text, src_p=None, keep_with_next=False):
        p = target_p.insert_paragraph_before()
        return _apply_source_formatting(p, line_text, src_p, keep_with_next)

    def _add_para_from_source(doc, line_text, src_p=None, keep_with_next=False):
        p = doc.add_paragraph()
        return _apply_source_formatting(p, line_text, src_p, keep_with_next)

    def _replace_in_paragraph(p, replacements, red_tags=None):
        full_text = p.text
        if not any(ph in full_text for ph in replacements):
            return False
        replaced = False
        r_tags = red_tags or set()
        for run in p.runs:
            for ph, val in replacements.items():
                if ph in run.text:
                    run.text = run.text.replace(ph, val)
                    if ph in r_tags:
                        try:
                            run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)
                            run.bold = True
                        except Exception:
                            pass
                    replaced = True
        if any(ph in p.text for ph in replacements):
            new_text = p.text
            has_red = False
            for ph, val in replacements.items():
                if ph in new_text:
                    new_text = new_text.replace(ph, val)
                    if ph in r_tags:
                        has_red = True
            if p.runs:
                p.runs[0].text = new_text
                if has_red:
                    try:
                        p.runs[0].font.color.rgb = docx.shared.RGBColor(255, 0, 0)
                        p.runs[0].bold = True
                    except Exception:
                        pass
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.text = new_text
            replaced = True
        return replaced

    def _copy_section_properties(src_sec, dst_sec):
        dst_sec.top_margin = src_sec.top_margin
        dst_sec.bottom_margin = src_sec.bottom_margin
        dst_sec.left_margin = src_sec.left_margin
        dst_sec.right_margin = src_sec.right_margin
        dst_sec.page_width = src_sec.page_width
        dst_sec.page_height = src_sec.page_height
        dst_sec.header_distance = src_sec.header_distance
        dst_sec.footer_distance = src_sec.footer_distance

    has_template = bool(template_docx_path.strip() and Path(template_docx_path).is_file())

    combined_doc = docx.Document()
    if has_template:
        tmpl_ref = docx.Document(template_docx_path)
        if tmpl_ref.sections:
            _copy_section_properties(tmpl_ref.sections[0], combined_doc.sections[0])
    else:
        for s in combined_doc.sections:
            s.top_margin = Cm(2.0)
            s.bottom_margin = Cm(2.0)
            s.left_margin = Cm(2.0)
            s.right_margin = Cm(1.5)

    def _slash_to_lines(text_val: str) -> str:
        """Конвертує слеші ' / ' або '/' (крім 'в/ч' та дат) у переноси рядків."""
        if not text_val:
            return ""
        t = re.sub(r"\b([вВ])\s*/\s*([чЧ])\b", r"\1_SLASH_TEMP_\2", str(text_val))
        t = re.sub(r"(\d)\s*/\s*(\d)", r"\1_NUMSLASH_TEMP_\2", t)
        t = re.sub(r"\s*/\s*", "\n", t)
        t = t.replace("_SLASH_TEMP_", "/")
        t = t.replace("_NUMSLASH_TEMP_", "/")
        return t

    signatory_title = _slash_to_lines(signatory_title)
    commander_signature_text = _slash_to_lines(commander_signature_text)
    executor_info = _slash_to_lines(executor_info)

    for extract_idx, (unit_code, unit_entry) in enumerate(units_data.items()):
        clean_unit = str(unit_code).replace("/", "-").replace("\\", "-").replace(" ", "_").strip()
        
        if isinstance(unit_entry, dict):
            header_lines = unit_entry.get("header_lines", [])
            items = unit_entry.get("items", [])
            recipient_to = str(unit_entry.get("recipient_to", "") or "").strip()
            destination_where = str(unit_entry.get("destination_where", "") or "").strip()
        else:
            header_lines = []
            items = [{"parent_heading": "", "text": str(x)} for x in unit_entry]
            recipient_to = ""
            destination_where = ""

        certification_block_text = ""
        if certify_extract:
            year_str = order_date[-4:] if len(order_date) >= 4 and order_date[-4:].isdigit() else "2026"
            certification_block_text = (
                f"Згідно з оригіналом:\n"
                f"{signatory_title}\n\n"
                f"{signatory_rank_name}\n"
                f"«____» ____________ {year_str} року"
            )

        if has_template:
            new_doc = docx.Document(template_docx_path)
            dest_val = destination_where.strip() if destination_where.strip() else "КУДИ"
            is_dest_manual = dest_val.upper() in ("КУДИ", "[КУДИ]")
            red_tags = {"{{куди}}", "{{Куди}}", "{{КУДИ}}"} if is_dest_manual else set()

            _all_mappings = {
                "{{кому}}": recipient_to,
                "{{Кому}}": recipient_to,
                "{{КОМУ}}": recipient_to,
                "{{куди}}": dest_val,
                "{{Куди}}": dest_val,
                "{{КУДИ}}": dest_val,
                "{{номер_наказу}}": order_number,
                "{{дата_наказу}}": order_date,
                "{{примірник}}": copy_number,
                "{{підпис_командира}}": commander_signature_text,
                "{{засвідчення}}": certification_block_text,
                "{{виконавець}}": executor_info,
            }
            replacements = {k: v for k, v in _all_mappings.items() if v.strip()}

            executor_replaced = False
            for p in list(new_doc.paragraphs):
                p_text = p.text
                if any(ph in p_text for ph in ["{{зміст}}", "{{content}}", "{{пункти}}"]):
                    if items:
                        printed_headings = set()
                        num_items = len(items)
                        for item_idx, item_data in enumerate(items):
                            heading = item_data.get("parent_heading", "") if isinstance(item_data, dict) else ""
                            text = item_data.get("text", "") if isinstance(item_data, dict) else str(item_data)
                            is_last_item = (item_idx == num_items - 1)

                            if heading and heading not in printed_headings:
                                src_h = source_para_map.get(heading.strip())
                                _insert_para_from_source(p, heading, src_p=src_h, keep_with_next=True)
                                printed_headings.add(heading)

                            lines = [line.rstrip() for line in text.splitlines() if line.strip()]
                            total_lines = len(lines)
                            for line_idx, line in enumerate(lines):
                                is_last_line_of_item = (line_idx == total_lines - 1)
                                kwn = not (is_last_item and is_last_line_of_item)
                                src_p = source_para_map.get(line.strip())
                                _insert_para_from_source(p, line, src_p=src_p, keep_with_next=kwn)
                    
                        p._element.getparent().remove(p._element)
                else:
                    if _replace_in_paragraph(p, replacements, red_tags=red_tags):
                        if "{{виконавець}}" in p_text:
                            executor_replaced = True

            if executor_info and not executor_replaced:
                _add_styled_p(new_doc, text="", space_after=8, first_indent=0)
                _add_styled_p(new_doc, text=executor_info, size_pt=8, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0, first_indent=0, keep_together=True)
        else:
            new_doc = docx.Document()
            sections = new_doc.sections
            for section in sections:
                section.top_margin = Cm(2.0)
                section.bottom_margin = Cm(2.0)
                section.left_margin = Cm(2.0)
                section.right_margin = Cm(1.5)

            if recipient_to:
                _add_styled_p(new_doc, text=recipient_to, size_pt=14, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2, first_indent=0, keep_with_next=True)
            if destination_where and destination_where.upper() not in ("КУДИ", "[КУДИ]"):
                _add_styled_p(new_doc, text=destination_where, size_pt=12, italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6, first_indent=0, keep_with_next=True)
            else:
                _add_styled_p(new_doc, text="КУДИ", size_pt=12, bold=True, color_rgb=(255, 0, 0), align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6, first_indent=0, keep_with_next=True)

            if header_lines:
                for idx, line in enumerate(header_lines):
                    align = WD_ALIGN_PARAGRAPH.CENTER if idx < len(header_lines) - 1 else WD_ALIGN_PARAGRAPH.LEFT
                    _add_styled_p(new_doc, text=line, size_pt=14, bold=(idx == 0), align=align, space_after=2, first_indent=0, keep_with_next=True)
            else:
                _add_styled_p(new_doc, text=f"НАКАЗ КОМАНДИРА ВІЙСЬКОВОЇ ЧАСТИНИ А0000", size_pt=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, first_indent=0, keep_with_next=True)
                _add_styled_p(new_doc, text=f"{order_date}               м. Львів               № {order_number}", size_pt=14, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, first_indent=0, keep_with_next=True)

            _add_styled_p(new_doc, text="", space_after=4, first_indent=0)
            _add_styled_p(new_doc, text=f"ВИТЯГ З НАКАЗУ", size_pt=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, first_indent=0, keep_with_next=True)
            if copy_number:
                _add_styled_p(new_doc, text=f"({copy_number})", size_pt=12, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8, first_indent=0, keep_with_next=True)

            printed_headings = set()
            num_items = len(items)
            for item_idx, item_data in enumerate(items):
                heading = item_data.get("parent_heading", "") if isinstance(item_data, dict) else ""
                text = item_data.get("text", "") if isinstance(item_data, dict) else str(item_data)
                is_last_item = (item_idx == num_items - 1)

                if heading and heading not in printed_headings:
                    src_h = source_para_map.get(heading.strip())
                    _add_para_from_source(new_doc, heading, src_p=src_h, keep_with_next=True)
                    printed_headings.add(heading)

                lines = [line.rstrip() for line in text.splitlines() if line.strip()]
                total_lines = len(lines)
                for line_idx, line in enumerate(lines):
                    is_last_line_of_item = (line_idx == total_lines - 1)
                    kwn = not (is_last_item and is_last_line_of_item)
                    src_p = source_para_map.get(line.strip())
                    _add_para_from_source(new_doc, line, src_p=src_p, keep_with_next=kwn)

            # Правило 5.3: рівно 2 порожні абзаци перед підписантом
            _add_styled_p(new_doc, text="", space_after=0, first_indent=0, keep_with_next=True)
            _add_styled_p(new_doc, text="", space_after=0, first_indent=0, keep_with_next=True)
            if commander_signature_text:
                for sig_line in commander_signature_text.splitlines():
                    _add_styled_p(new_doc, text=sig_line, size_pt=14, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=3, first_indent=0, keep_with_next=True, keep_together=True)

            if certify_extract:
                _add_styled_p(new_doc, text="", space_after=12, first_indent=0, keep_with_next=True)
                _add_styled_p(new_doc, text="Згідно з оригіналом:", size_pt=14, bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4, first_indent=0, keep_with_next=True, keep_together=True)
                _add_styled_p(new_doc, text=signatory_title, size_pt=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4, first_indent=0, keep_with_next=True, keep_together=True)
                _add_styled_p(new_doc, text="", space_after=8, first_indent=0, keep_with_next=True)
                _add_styled_p(new_doc, text=f"{signatory_rank_name}", size_pt=14, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=4, first_indent=0, keep_together=True)
                _add_styled_p(new_doc, text=f"«____» ____________ {order_date[-4:] if len(order_date)>=4 and order_date[-4:].isdigit() else '2026'} року", size_pt=12, italic=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4, first_indent=0)

            if executor_info:
                _add_styled_p(new_doc, text="", space_after=8, first_indent=0)
                _add_styled_p(new_doc, text=executor_info, size_pt=8, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0, first_indent=0, keep_together=True)

        out_name = f"Витяг_{clean_unit}_№{order_number}.docx"
        out_file = out_dir / out_name
        if save_individual_files:
            new_doc.save(out_file)
            created_paths.append(str(out_file))

        table_rows.append((unit_code, len(items), out_name, str(out_file)))

        # Додаємо елементи витягу у спільний документ для двостороннього друку
        if extract_idx > 0:
            sec = combined_doc.add_section(WD_SECTION.ODD_PAGE)
            sec.start_type = WD_SECTION.ODD_PAGE
            if has_template:
                _copy_section_properties(tmpl_ref.sections[0], sec)
            else:
                sec.top_margin = Cm(2.0)
                sec.bottom_margin = Cm(2.0)
                sec.left_margin = Cm(2.0)
                sec.right_margin = Cm(1.5)

        for elem in new_doc.element.body:
            if elem.tag.endswith("sectPr"):
                continue
            elem_copy = copy.deepcopy(elem)
            for child in list(elem_copy.iter()):
                if child.tag.endswith("sectPr"):
                    p_parent = child.getparent()
                    if p_parent is not None:
                        p_parent.remove(child)
            combined_doc.element.body.append(elem_copy)

    clean_order_num = str(order_number).replace("/", "-").replace("\\", "-").strip()
    combined_path = out_dir / f"Всі_витяги_наказ_№{clean_order_num}.docx"
    combined_doc.save(combined_path)

    final_paths = created_paths if save_individual_files else [str(combined_path)]

    table = DataTable(
        ("Військова частина", "Кількість пунктів", "Файл витягу", "Повний шлях"),
        tuple(table_rows),
    )

    return {
        "summary": f"Згенеровано єдиний файл з усіма витягами ({len(units_data)} адресатів) для двостороннього друку: {combined_path.name}",
        "count": len(final_paths),
        "details": table,
        "paths": final_paths,
        "combined_path": str(combined_path),
    }
