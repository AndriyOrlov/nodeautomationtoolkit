"""Фікстури для наскрізного прогону ПОВІДОМЛЕНЬ (модуль 3).

Усе вигадане, усе на C:. Особливості, заради яких ця фікстура існує:

* у наказі шрифт заданий СТИЛЕМ `Normal` (Times New Roman 14), а не окремо
  в кожному рядку — саме так Word не записує ознаку шрифту в абзац, і
  скопійований `FormattedText` успадковує `Normal` ШАБЛОНА (правило 5.2.1);
* у шаблонах повідомлень `Normal` навмисно ЧУЖИЙ (Calibri 11) — якщо шрифт
  не переносити явно, зміст повідомлення виявиться не Times New Roman;
* у наказі є частина з почесним найменуванням у лапках і корпус,
  зазначений у стовпці D скороченням, — перевірка правил 9.5.
"""
import os
import sys

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from openpyxl import Workbook

OUT = sys.argv[1]
os.makedirs(OUT, exist_ok=True)
C = WD_ALIGN_PARAGRAPH.CENTER
J = WD_ALIGN_PARAGRAPH.JUSTIFY
R = WD_ALIGN_PARAGRAPH.RIGHT


def para(doc, text="", align=None, indent=None, left=None, bold=False):
    """Абзац БЕЗ явного шрифту: шрифт успадковується від стилю `Normal`."""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if indent is not None:
        p.paragraph_format.first_line_indent = Cm(indent)
    if left is not None:
        p.paragraph_format.left_indent = Cm(left)
        p.paragraph_format.first_line_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    return p


def bio(doc, *lines):
    p = para(doc, lines[0], left=8.0)
    for extra in lines[1:]:
        p.add_run().add_break()
        p.add_run(extra)
    return p


def item(doc, text):
    return para(doc, text, align=J, indent=1.25)


# ── Наказ: шрифт лише у стилі Normal ─────────────────────────────────────────
order = docx.Document()
normal = order.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(14)
s = order.sections[0]
s.left_margin, s.right_margin = Cm(2.0), Cm(1.0)
s.top_margin, s.bottom_margin = Cm(1.0), Cm(1.0)

para(order, "МІНІСТЕРСТВО ОБОРОНИ УКРАЇНИ", C, bold=True)
para(order)
para(order, "НАКАЗ", C, bold=True)
para(order, "КОМАНДИРА ВІЙСЬКОВОЇ ЧАСТИНИ А0001", C, bold=True)
para(order)
para(order, "(по особовому складу)", C)
para(order)
para(order, "«15» серпня 2026 року          м. Тестове          № 557", bold=True)
para(order)

para(order, "§ 1", C)
para(order)
para(order, "Відповідно до пункту 1 Положення про проходження військової служби "
            "з нижчепойменованими військовослужбовцями ПРИЗНАЧИТИ:", align=J)
para(order)

# 1 — почесне найменування в лапках + корпус, названий і в стовпці D, і в тексті
item(order, "1. Лейтенанта ТЕСТЕНКА Андрія Андрійовича, командира тестового "
            "взводу тестової роти тестового батальйону "
            "77 окремої тестової бригади «Тестовий Яр» 51 армійського корпусу, "
            "КОМАНДИРОМ ТЕСТОВОГО ВЗВОДУ ТЕСТОВОЇ РОТИ ТЕСТОВОГО БАТАЛЬЙОНУ "
            "88 ОКРЕМОЇ ЗРАЗКОВОЇ БРИГАДИ.")
bio(order, "1990 р. н., освіта: ТВІ у 2012 р.,", "у ЗС - із 08.2008.")
bio(order, "1234567890.")
para(order)

# 1.1 — трирівневий ланцюг: батальйон → бригада → корпус. Назву РВЕ мʼякий
# перенос (Shift+Enter), як у справжніх наказах (правило 4.2.8), а в стовпці D
# батальйону стоїть КОРПУС — саме на цьому ланцюг ставав «А1111 А3333 А2222».
p = item(order, "1.1. Лейтенанта ЛАНЦЮЖЕНКА Дмитра Дмитровича, командира ")
p.add_run("тестового взводу 66 окремого тестового батальйону")
p.add_run().add_break()
p.add_run("77 окремої тестової бригади «Тестовий Яр» 51 армійського корпусу, ")
p.add_run("командиром тестового взводу цієї самої бригади.")
bio(order, "1992 р. н., освіта: ТВІ у 2014 р.,", "у ЗС - із 08.2010.")
bio(order, "4567890123.")
para(order)

# 1.2 — ТЦК: «звідки» малими, «КУДИ» ВЕЛИКИМИ. Один збіг словника тягнувся від
# «…Тестової області» аж до «…СОЦІАЛЬНОЇ ПІДТРИМКИ» у ВЕЛИКІЙ половині і з'їдав
# «– НАЧАЛЬНИКОМ ГРУПИ …»: 139 символів наказу зникали (розд. 4.2.9 та 9.5.5).
item(order, "1.2. Капітана ТЦКАЄНКА Сергія Сергійовича, офіцера відділення "
            "рекрутингу та комплектування Прикладного районного "
            "територіального центру комплектування та соціальної підтримки "
            "Тестової області – НАЧАЛЬНИКОМ ГРУПИ ПСИХОЛОГІЧНОЇ ПІДТРИМКИ "
            "ПЕРСОНАЛУ ПРИКЛАДНОГО РАЙОННОГО ТЕРИТОРІАЛЬНОГО ЦЕНТРУ "
            "КОМПЛЕКТУВАННЯ ТА СОЦІАЛЬНОЇ ПІДТРИМКИ ЦІЄЇ САМОЇ ОБЛАСТІ.")
bio(order, "1988 р. н., освіта: ТВІ у 2010 р.,", "у ЗС - із 08.2006.")
bio(order, "5678901234.")
para(order)

# 2 — зворот «цієї самої бригади» (правило 9.6) та відмінок (правило 9.5)
item(order, "2. Майора ПРИКЛАДЕНКА Богдана Богдановича, начальника служби "
            "77 окремої тестової бригади «Тестовий Яр», начальником служби "
            "цієї самої бригади.")
bio(order, "1985 р. н., освіта: ТВІ у 2007 р.,", "у ЗС - із 09.2003.")
bio(order, "2345678901.")
para(order)
para(order)

para(order, "§ 2", C)
para(order)
para(order, "Відповідно до підпункту «б» пункту 2 ЗВІЛЬНИТИ З ВІЙСЬКОВОЇ СЛУЖБИ:", align=J)
para(order)

# 3 — ТЦК у біографічному рядку НЕ стає адресатом (правило 9.9)
item(order, "3. Старшого лейтенанта ЗРАЗКОВА Василя Васильовича, офіцера "
            "88 окремої зразкової бригади.")
bio(order, "1995 р. н., освіта: ТВІ у 2017 р.,", "у ЗС - із 08.2013.")
bio(order, "3456789012.")
bio(order, "Підлягає направленню на військовий облік до Тестового",
    "районного територіального центру комплектування та соціальної підтримки.")
para(order)
para(order)

para(order, "Командир військової частини А0001")
para(order)
para(order, "полковник                                        Петро ТЕСТОВИЙ")

order_path = os.path.join(OUT, "Наказ № 557 від 15.08.2026.docx")
order.save(order_path)

# ── Словник: стовпець D містить СКОРОЧЕННЯ корпусу (звичайний випадок) ───────
wb = Workbook()
ws = wb.active
ws.append(["Найменування", "Шифр", "Скорочення", "Корпус", "Кому", "Куди"])
for row in [
    ("51 армійський корпус", "А0051", "51 АК", "",
     "Командиру військової частини А0051", "м. Корпусне"),
    ("66 окремий тестовий батальйон", "А0066", "66 отб", "51 АК",
     "Командиру військової частини А0066", "м. Нульове"),
    ("77 окрема тестова бригада", "А0077", "77 отбр", "51 АК",
     "Командиру військової частини А0077", "м. Перше"),
    ("88 окрема зразкова бригада", "А0088", "88 озбр", "",
     "Командиру військової частини А0088", "м. Друге"),
    ("Тестовий обласний територіальний центр комплектування та соціальної підтримки",
     "", "Тестовий ОТЦК та СП", "",
     "Начальнику Тестового обласного ТЦК та СП", "м. Тестове"),
]:
    ws.append(list(row))
wb.save(os.path.join(OUT, "mapping.xlsx"))


def message_template(path, with_content):
    """Шаблон повідомлення з НАВМИСНО чужим стилем `Normal` (Calibri 11)."""
    tpl = docx.Document()
    tpl_normal = tpl.styles["Normal"]
    tpl_normal.font.name = "Calibri"
    tpl_normal.font.size = Pt(11)
    sect = tpl.sections[0]
    sect.left_margin, sect.right_margin = Cm(2.0), Cm(1.0)
    sect.top_margin, sect.bottom_margin = Cm(1.0), Cm(1.0)

    table = tpl.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "{{тцк чі вч}}"
    table.cell(0, 1).text = "{{куди}}"
    for row_index in range(1, 4):
        table.cell(row_index, 0).text = "{{кому_список}}"
        table.cell(row_index, 1).text = ""

    para(tpl)
    para(tpl, "Про прийняття кадрових рішень", align=C)
    para(tpl)
    para(tpl, "Надсилаємо витяг з наказу {{номер_наказу}} від {{дата_наказу}}.", align=J)
    para(tpl)
    if with_content:
        para(tpl, "{{зміст_шифр}}")
        para(tpl)
    para(tpl, "{{виконавець}}")
    para(tpl, "ВІДКРИТА ІНФОРМАЦІЯ")
    tpl.save(path)


message_template(os.path.join(OUT, "message_cover.docx"), with_content=False)
message_template(os.path.join(OUT, "message_content.docx"), with_content=True)

print("готово:", order_path)
