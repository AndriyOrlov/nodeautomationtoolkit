"""Багатосторінковий тестовий наказ: перевірка макета на межі сторінок.

Відрізняється від `make_fixtures.py` тим, що на ОДНУ частину припадає стільки
пунктів, що витяг гарантовано не вміщується на одну сторінку. Саме тут працює
логіка, якої не видно на короткому наказі:

* підбір міжрядкового інтервалу для БАГАТОСТОРІНКОВОГО витягу (скільки пунктів
  сідає на першу сторінку);
* перевірка макета `layout_issues` (вона запускається лише коли сторінок > 1);
* зчеплення шапки (§) з першим пунктом розділу через порожній абзац між ними —
  саме там Word рвав сторінку й шапка лишалась унизу сама;
* вирівнювання аркушів для друку «2 на 1».

Кількість пунктів підібрано так, щоб § 2 опинявся приблизно на межі сторінок.

Усе вигадане, усе на C:.

    python scripts/e2e_extracts/make_fixtures_multipage.py <тека>
    python scripts/e2e_extracts/run_extracts_e2e.py <тека> .
"""
import os
import sys

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from openpyxl import Workbook

OUT = sys.argv[1]
# Скільки пунктів стоїть ДО § 2 — цим параметром § заганяється на межу сторінки.
ITEMS_BEFORE_SECOND = int(sys.argv[2]) if len(sys.argv) > 2 else 8
os.makedirs(OUT, exist_ok=True)
C = WD_ALIGN_PARAGRAPH.CENTER
J = WD_ALIGN_PARAGRAPH.JUSTIFY
R = WD_ALIGN_PARAGRAPH.RIGHT


def para(doc, text="", align=None, indent=None, bold=False, size=14, left=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if indent is not None:
        p.paragraph_format.first_line_indent = Cm(indent)
    if left is not None:
        p.paragraph_format.left_indent = Cm(left)
        p.paragraph_format.first_line_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    return p


def bio(doc, *lines):
    p = para(doc, lines[0], left=8.0)
    for extra in lines[1:]:
        p.add_run().add_break()
        run = p.add_run(extra)
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
    return p


def item(doc, text):
    return para(doc, text, align=J, indent=1.25)


NAMES = [
    ("ТЕСТЕНКА Андрія Андрійовича", "капітана"),
    ("ПРИКЛАДЕНКА Богдана Богдановича", "майора"),
    ("ЗРАЗКОВА Василя Васильовича", "старшого лейтенанта"),
    ("ДОСЛІДЕНКА Григорія Григоровича", "капітана"),
    ("ПЕРЕВІРЕНКА Дмитра Дмитровича", "старшого солдата"),
    ("МАКЕТЕНКА Євгена Євгеновича", "сержанта"),
    ("ШАБЛОНЕНКА Жоржа Жоржовича", "капітана"),
    ("ВЕРСТАЄНКА Захара Захаровича", "майора"),
    ("СТОРІНКОВА Івана Івановича", "лейтенанта"),
    ("АРКУШЕНКА Кирила Кириловича", "солдата"),
    ("ІНТЕРВАЛЕНКА Леоніда Леонідовича", "прапорщика"),
    ("ПАГІНАЦЬКОГО Миколу Миколайовича", "капітана"),
]


def add_item(doc, number, unit_phrase):
    name, rank = NAMES[(number - 1) % len(NAMES)]
    item(
        doc,
        f"{number}. {rank.capitalize()} {name}, командира взводу {unit_phrase}, "
        f"призначити на посаду командира роти {unit_phrase} з випробувальним "
        f"строком, установленим законодавством України.",
    )
    bio(doc, f"199{number % 10} р. н., освіта: ТВІ у 201{number % 10} р.,", "у ЗС - із 08.2010.")
    bio(doc, f"12345678{number:02d}.")
    para(doc)


order = docx.Document()
s = order.sections[0]
s.left_margin, s.right_margin = Cm(2.0), Cm(1.0)
s.top_margin, s.bottom_margin = Cm(1.0), Cm(1.0)

para(order, "МІНІСТЕРСТВО ОБОРОНИ УКРАЇНИ", C, bold=True)
para(order)
para(order, "НАКАЗ", C, bold=True, size=22)
para(order, "КОМАНДИРА ВІЙСЬКОВОЇ ЧАСТИНИ А0001", C, bold=True)
para(order)
para(order, "(по особовому складу)", C)
para(order)
para(order, "«15» серпня 2026 року          м. Тестове          № 557", bold=True)
para(order)

para(order, "§ 1", C)
para(order)
para(order, "Відповідно до пункту 1 Положення про проходження військової служби "
            "з нижчепойменованими військовослужбовцями ПРИЗНАЧИТИ:", align=J)
para(order)

# Вісім пунктів на ОДНУ частину — витяг для неї стане багатосторінковим.
for number in range(1, ITEMS_BEFORE_SECOND + 1):
    add_item(order, number, "3 окремого тестового загону")

# § 2 починається приблизно на межі сторінок — саме тут перевіряється, чи не
# лишиться шапка внизу сторінки без свого першого пункту.
para(order, "§ 2", C)
para(order)
para(order, "Відповідно до підпункту «б» пункту 2 ЗВІЛЬНИТИ З ВІЙСЬКОВОЇ СЛУЖБИ:", align=J)
para(order)

for number in range(ITEMS_BEFORE_SECOND + 1, ITEMS_BEFORE_SECOND + 5):
    add_item(order, number, "3 окремого тестового загону")

para(order, "Командир військової частини А0001")
para(order)
para(order, "полковник                                        Петро ТЕСТОВИЙ")
para(order)

para(order, "Розрахунок розсилки витягів із наказу:", C)
para(order, "1. Військова частина А0003\tпп. 1-12.")
para(order, "Надр. 1 прим.")

order_path = os.path.join(OUT, "Наказ № 557 від 15.08.2026.docx")
order.save(order_path)

wb = Workbook()
ws = wb.active
ws.append(["Найменування", "Шифр", "Скорочення", "Корпус", "Кому", "Куди"])
for row in [
    ("3 окремий тестовий загін", "А0003", "3 отз", "",
     "Командиру військової частини А0003", "м. Третє"),
    ("4 окремий тестовий полк", "А0004", "4 отп", "",
     "Командиру військової частини А0004", "м. Четверте"),
]:
    ws.append(list(row))
wb.save(os.path.join(OUT, "mapping.xlsx"))

tpl = docx.Document()
s = tpl.sections[0]
s.left_margin, s.right_margin = Cm(2.0), Cm(1.0)
s.top_margin, s.bottom_margin = Cm(1.0), Cm(1.0)
normal = tpl.styles["Normal"]
normal.paragraph_format.first_line_indent = Cm(1.25)
normal.paragraph_format.alignment = J

para(tpl, "{{кому}}", R)
para(tpl, "{{куди}}", R)
para(tpl)
para(tpl, "ВИТЯГ З НАКАЗУ", C, bold=True)
para(tpl, "КОМАНДИРА ВІЙСЬКОВОЇ ЧАСТИНИ А0001", C, bold=True)
para(tpl)
para(tpl, "{{дата_наказу}}          м. Тестове          {{номер_наказу}}")
para(tpl)
para(tpl, "{{зміст}}")
para(tpl)
para(tpl, "{{засвідчення}}")
para(tpl)
p = tpl.add_paragraph()
run = p.add_run("{{виконавець}}")
run.font.name = "Times New Roman"
run.font.size = Pt(8)
tpl.save(os.path.join(OUT, "template.docx"))

print("готово:", order_path)
