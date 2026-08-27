"""Розширений тестовий наказ: корпус, підпорядкована частина, ТЦК,
внутрішнє переміщення, розрахунок розсилки У ВИГЛЯДІ ТАБЛИЦІ.

Усе вигадане, усе на C:. Структура — за додатками 41/43/44 (AGENT.md розд. 12).
"""
import os
import sys

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from openpyxl import Workbook

OUT = sys.argv[1]
os.makedirs(OUT, exist_ok=True)
C = WD_ALIGN_PARAGRAPH.CENTER
J = WD_ALIGN_PARAGRAPH.JUSTIFY


def para(doc, text="", align=None, indent=None, bold=False, size=14, left=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if indent is not None:
        p.paragraph_format.first_line_indent = Cm(indent)
    if left is not None:
        p.paragraph_format.left_indent = Cm(left)
        p.paragraph_format.first_line_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    return p


def bio(doc, *lines):
    p = para(doc, lines[0], left=8.0)
    for extra in lines[1:]:
        p.add_run().add_break()
        run = p.add_run(extra)
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
    return p


def item(doc, text):
    return para(doc, text, align=J, indent=1.25)


order = docx.Document()
s = order.sections[0]
s.left_margin, s.right_margin = Cm(2.0), Cm(1.0)
s.top_margin, s.bottom_margin = Cm(1.0), Cm(1.0)

para(order, "МІНІСТЕРСТВО ОБОРОНИ УКРАЇНИ", C, bold=True)
para(order)
para(order, "НАКАЗ", C, bold=True, size=22)
para(order, "КОМАНДИРА ВІЙСЬКОВОЇ ЧАСТИНИ А0001", C, bold=True)
para(order)
para(order, "(по особовому складу)", C)
para(order)
para(order, "«15» серпня 2026 року          м. Тестове          № 555", bold=True)
para(order)

para(order, "§ 1", C)
para(order)
para(order, "Відповідно до пункту 1 Положення про проходження військової служби "
            "з нижчепойменованими військовослужбовцями ПРИЗНАЧИТИ:", align=J)
para(order)

# 1 — підпорядкована корпусу частина: витяг має піти НА КОРПУС
item(order, "1. Капітана ТЕСТЕНКА Андрія Андрійовича, командира роти "
            "1 окремого тестового батальйону 90 армійського корпусу, "
            "командиром роти 2 окремого тестового полку 90 армійського корпусу.")
bio(order, "1990 р. н., освіта: ТВІ у 2012 р.,", "у ЗС - із 08.2008.")
bio(order, "1234567890.")
para(order)

# 2 — самостійна частина, БЕЗ порожнього перед біографією (перевірка нормалізації)
item(order, "2. Майора ПРИКЛАДЕНКА Богдана Богдановича, начальника служби "
            "3 окремого тестового загону, начальником служби "
            "цього самого загону.")
bio(order, "1985 р. н., освіта: ТВІ у 2007 р.,", "у ЗС - із 09.2003.")
bio(order, "2345678901.")
bio(order, "Підстава: клопотання командира військової частини.")
para(order)
para(order)

para(order, "§ 2", C)
para(order)
para(order, "Відповідно до підпункту «б» пункту 2 ЗВІЛЬНИТИ З ВІЙСЬКОВОЇ СЛУЖБИ:", align=J)
para(order)

# 3 — ТЦК у біографічному рядку НЕ має стати адресатом (правило 4.2)
item(order, "3. Старшого лейтенанта ЗРАЗКОВА Василя Васильовича, офіцера "
            "4 окремого тестового полку.")
para(order, left=8.0)
bio(order, "1995 р. н., освіта: ТВІ у 2017 р.,", "у ЗС - із 08.2013.")
bio(order, "3456789012.")
bio(order, "Підлягає направленню на військовий облік до Тестового",
    "районного територіального центру комплектування та соціальної підтримки.")
para(order)
para(order)

para(order, "Командир військової частини А0001")
para(order)
para(order, "полковник                                        Петро ТЕСТОВИЙ")
para(order)

# ── Зворот останнього аркуша: розрахунок У ТАБЛИЦІ + другий блок ────────────
para(order, "Розрахунок розсилки витягів із наказу:", C)
table = order.add_table(rows=3, cols=4)
table.style = "Table Grid"
data = [
    ("1. Військова частина А0090", "п. 1.", "3. Військова частина А0004", "п. 3."),
    ("2. Військова частина А0003", "п. 2.", "", ""),
    ("Усього:", "3", "прим.", ""),
]
for row_index, row in enumerate(data):
    for col_index, value in enumerate(row):
        cell = table.cell(row_index, col_index)
        cell.text = value
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(14)

para(order, "Розрахунок розсилки електронних повідомлень:", C)
para(order, "1. Військова частина А0090\tп. 1.\t2. Військова частина А0003\tп. 2.")
para(order, "Надр. 2 прим.")
para(order, "Прим. № 1 — перший адресат.")

order_path = os.path.join(OUT, "Наказ № 556 від 15.08.2026.docx")
order.save(order_path)

# ── Словник: корпус окремим рядком + підпорядковані частини + ТЦК ────────────
wb = Workbook()
ws = wb.active
ws.append(["Найменування", "Шифр", "Скорочення", "Корпус", "Кому", "Куди"])
for row in [
    ("90 армійський корпус", "А0090", "90 АК", "",
     "Командиру військової частини А0090", "м. Корпусне"),
    ("1 окремий тестовий батальйон", "А0002", "1 отб", "90 АК",
     "Командиру військової частини А0002", "м. Перше"),
    ("2 окремий тестовий полк", "А0091", "2 отп", "90 АК",
     "Командиру військової частини А0091", "м. Друге"),
    ("3 окремий тестовий загін", "А0003", "3 отз", "",
     "Командиру військової частини А0003", "м. Третє"),
    ("4 окремий тестовий полк", "А0004", "4 отп", "",
     "Командиру військової частини А0004", "м. Четверте"),
    ("Тестовий обласний територіальний центр комплектування та соціальної підтримки",
     "", "Тестовий ОТЦК та СП", "",
     "Начальнику Тестового обласного ТЦК та СП", "м. Тестове"),
]:
    ws.append(list(row))
wb.save(os.path.join(OUT, "mapping.xlsx"))

# ── Шаблон із «ворожим» стилем Normal (перевірка незалежності від шаблону) ───
tpl = docx.Document()
s = tpl.sections[0]
s.left_margin, s.right_margin = Cm(2.0), Cm(1.0)
s.top_margin, s.bottom_margin = Cm(1.0), Cm(1.0)
normal = tpl.styles["Normal"]
normal.paragraph_format.first_line_indent = Cm(1.25)
normal.paragraph_format.alignment = J

R = WD_ALIGN_PARAGRAPH.RIGHT
para(tpl, "{{кому}}", R)
para(tpl, "{{куди}}", R)
para(tpl)
para(tpl, "ВИТЯГ З НАКАЗУ", C, bold=True)
para(tpl, "КОМАНДИРА ВІЙСЬКОВОЇ ЧАСТИНИ А0001", C, bold=True)
para(tpl)
para(tpl, "{{дата_наказу}}          м. Тестове          {{номер_наказу}}")
para(tpl)
para(tpl, "{{зміст}}")
para(tpl)
para(tpl, "{{засвідчення}}")
para(tpl)
p = tpl.add_paragraph()
run = p.add_run("{{виконавець}}")
run.font.name = "Times New Roman"
run.font.size = Pt(8)
tpl.save(os.path.join(OUT, "template.docx"))

print("готово:", order_path)
