"""Наказ за формою зі скріншота користувача: пункт ЗАРАХУВАТИ у розпорядження,
біоблок, Підстава, підписний блок із багаторядковою посадою."""
import os, sys
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from openpyxl import Workbook

OUT = sys.argv[1]; os.makedirs(OUT, exist_ok=True)
C = WD_ALIGN_PARAGRAPH.CENTER; J = WD_ALIGN_PARAGRAPH.JUSTIFY

def para(doc, text="", align=None, indent=None, bold=False, size=14, left=None):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    if indent is not None: p.paragraph_format.first_line_indent = Cm(indent)
    if left is not None:
        p.paragraph_format.left_indent = Cm(left)
        p.paragraph_format.first_line_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text); r.font.name = "Times New Roman"; r.font.size = Pt(size); r.bold = bold
    return p

def bio(doc, *lines):
    p = para(doc, lines[0], left=8.0)
    for extra in lines[1:]:
        p.add_run().add_break()
        r = p.add_run(extra); r.font.name = "Times New Roman"; r.font.size = Pt(14)
    return p

d = docx.Document()
s = d.sections[0]
s.left_margin, s.right_margin = Cm(2.0), Cm(1.0); s.top_margin, s.bottom_margin = Cm(1.0), Cm(1.0)

para(d, "МІНІСТЕРСТВО ОБОРОНИ УКРАЇНИ", C, bold=True); para(d)
para(d, "НАКАЗ", C, bold=True, size=22)
para(d, "КОМАНДУВАЧА ВІЙСЬК ОПЕРАТИВНОГО КОМАНДУВАННЯ «ТЕСТ»", C, bold=True); para(d)
para(d, "(по особовому складу)", C); para(d)
para(d, "«15» серпня 2026 року          м. Тестове          № 555", bold=True); para(d)

para(d, "§ 1", C); para(d)
para(d, "Відповідно до пунктів 1 та 2 Положення про проходження громадянами України "
        "військової служби у Збройних Силах України нижчепойменованих осіб офіцерського "
        "складу ЗВІЛЬНИТИ з займаних посад:", align=J)
para(d)
para(d, "1. Полковника ТЕСТЕНКА Андрія Андрійовича, командира роти, звільнити із "
        "займаної посади і ЗАРАХУВАТИ у розпорядження командувача військ оперативного "
        "командування «Тест». На час перебування у розпорядженні командувача військ "
        "оперативного командування «Тест» залишається на всіх видах забезпечення та "
        "у списках 267 окремої механізованої бригади.", align=J, indent=1.25)
bio(d, "1971 р. н.", "1234567890.")
bio(d, "Підстава: план переміщення командира", "військової частини А0267 від 06.08.2026", "№ 12.")
para(d); para(d)

# Підписний блок: посада у ТРИ рядки, розсунута порожніми абзацами
para(d, "Тимчасово виконуючий обов'язки")
para(d, "командувача військ")
para(d, "оперативного командування «Тест»")
para(d); para(d)
para(d, "полковник                                        Іван ПЕТРЕНКО")
para(d)

para(d, "Розрахунок розсилки витягів із наказу:", C)
para(d, "1. Військова частина А0777\tп. 1.")
para(d, "Надр. 2 прим.")

order = os.path.join(OUT, "Наказ № 555 від 15.08.2026.docx"); d.save(order)

wb = Workbook(); ws = wb.active
ws.append(["Найменування","Шифр","Скорочення","Корпус","Кому","Куди"])
ws.append(["оперативне командування «Тест»","А0777","ОК Тест","","Командиру військової частини А0777","м. Тестове"])
ws.append(["267 окрема механізована бригада","А0267","267 омбр","ОК Тест","Командиру військової частини А0267","м. Бригадне"])
wb.save(os.path.join(OUT,"mapping.xlsx"))

t = docx.Document()
s = t.sections[0]
s.left_margin, s.right_margin = Cm(2.0), Cm(1.0); s.top_margin, s.bottom_margin = Cm(1.0), Cm(1.0)
R = WD_ALIGN_PARAGRAPH.RIGHT
para(t,"{{кому}}",R); para(t,"{{куди}}",R); para(t)
para(t,"ВИТЯГ З НАКАЗУ",C,bold=True); para(t)
para(t,"{{дата_наказу}}          м. Тестове          {{номер_наказу}}"); para(t)
para(t,"{{зміст}}"); para(t); para(t,"{{засвідчення}}"); para(t)
p=t.add_paragraph(); r=p.add_run("{{виконавець}}"); r.font.name="Times New Roman"; r.font.size=Pt(8)
t.save(os.path.join(OUT,"template.docx"))
print("готово:", order)
