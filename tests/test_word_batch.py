from docx import Document

from nodeautomationtoolkit.builtin_nodes.word_batch import (
    _normalize_page_selector,
    batch_delete_page,
    batch_fill_placeholder,
    batch_insert_page,
    batch_keep_items_together,
    batch_replace_text,
    create_document_batch,
    save_document_batch,
)
from nodeautomationtoolkit.core.batch_types import WordDocumentBatch


def make_source(path):
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Початок {{ТЕ")
    bold_run = paragraph.add_run("КСТ}} кінець")
    bold_run.bold = True
    document.save(path)


def test_builds_batch_from_groups_and_expands_fields(tmp_path):
    source = tmp_path / "order.docx"
    make_source(source)
    batch = create_document_batch(
        str(source),
        groups={"АЛЬФА": "Пункт для Альфи", "БРАВО": "Пункт для Браво"},
    )
    batch = batch_fill_placeholder(batch, "{{ТЕКСТ}}", "{{content}}")
    batch = batch_replace_text(batch, "Початок", "Витяг {{name}}")
    batch = batch_keep_items_together(batch)

    assert isinstance(batch, WordDocumentBatch)
    assert len(batch.variants) == 2
    assert len(batch.operations) == 3


def test_builds_batch_from_cipher_document_fields(tmp_path):
    source = tmp_path / "order.docx"
    make_source(source)
    batch = create_document_batch(
        str(source),
        groups={
            "Ш-01": {
                "name": "Ш-01",
                "content": "1. Пункт",
                "destination": "Відділ 1",
            }
        },
    )

    assert batch.variants[0].values()["destination"] == "Відділ 1"


def test_saves_all_text_only_variants_without_word_com(tmp_path):
    source = tmp_path / "order.docx"
    output = tmp_path / "result"
    make_source(source)
    batch = create_document_batch(
        str(source),
        groups={"АЛЬФА": "Пункт для Альфи", "БРАВО": "Пункт для Браво"},
    )
    batch = batch_fill_placeholder(batch, "{{ТЕКСТ}}", "{{content}}")

    result = save_document_batch(batch, str(output))

    assert result["count"] == 2
    alpha = Document(output / "АЛЬФА.docx")
    bravo = Document(output / "БРАВО.docx")
    assert alpha.paragraphs[0].text == "Початок Пункт для Альфи кінець"
    assert bravo.paragraphs[0].text == "Початок Пункт для Браво кінець"
    assert alpha.paragraphs[0].runs[-1].bold is True


def test_adds_delete_page_operation_with_first_last_or_number(tmp_path):
    source = tmp_path / "order.docx"
    make_source(source)
    batch = create_document_batch(str(source))

    first = batch_delete_page(batch, "перша")
    last = batch_delete_page(batch, "остання")
    numbered = batch_delete_page(batch, "3")

    assert first.operations[-1].options()["page"] == "first"
    assert last.operations[-1].options()["page"] == "last"
    assert numbered.operations[-1].options()["page"] == 3


def test_adds_blank_or_template_page_operation(tmp_path):
    source = tmp_path / "order.docx"
    page_template = tmp_path / "page.docx"
    make_source(source)
    make_source(page_template)
    batch = create_document_batch(str(source))

    blank = batch_insert_page(batch, "перша")
    from_template = batch_insert_page(batch, "2", str(page_template))

    assert blank.operations[-1].options() == {
        "page": "first",
        "page_docx": "",
    }
    assert from_template.operations[-1].options() == {
        "page": 2,
        "page_docx": str(page_template.resolve()),
    }


def test_page_selector_rejects_unknown_or_non_positive_values():
    for value in ("нульова", "0", -2):
        try:
            _normalize_page_selector(value)
        except ValueError:
            pass
        else:
            raise AssertionError(f"Очікувалась помилка для значення {value!r}")
