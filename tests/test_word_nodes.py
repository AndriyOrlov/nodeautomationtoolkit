from pathlib import Path

from docx import Document

from nodeautomationtoolkit.builtin_nodes.word import (
    create_docx,
    filter_paragraphs,
    paragraphs_to_text,
    read_docx,
    save_selected_paragraphs,
)
from nodeautomationtoolkit.core.registry import NodeRegistry
from nodeautomationtoolkit.core.word_types import WordDocument, WordParagraphs


def make_source(path: Path) -> None:
    document = Document()
    document.add_heading("Тестовий наказ", level=1)
    document.add_paragraph("Перший пункт для Альфи.")
    document.add_paragraph("Другий пункт для Браво.")
    document.save(path)


def test_reads_and_filters_specialized_word_types(tmp_path: Path):
    source_path = tmp_path / "order.docx"
    make_source(source_path)

    result = read_docx(str(source_path))
    selected = filter_paragraphs(result["paragraphs"], "альфи")

    assert isinstance(result["document"], WordDocument)
    assert isinstance(result["paragraphs"], WordParagraphs)
    assert result["file_name"] == "order.docx"
    assert len(selected) == 1
    assert paragraphs_to_text(selected) == "Перший пункт для Альфи."


def test_creates_docx_and_preserves_selected_paragraph_formatting(tmp_path: Path):
    source_path = tmp_path / "source.docx"
    selected_path = tmp_path / "selected.docx"
    created_path = tmp_path / "created.docx"
    make_source(source_path)
    result = read_docx(str(source_path))
    selected = filter_paragraphs(result["paragraphs"], "Другий пункт")

    saved = save_selected_paragraphs(
        result["document"],
        selected,
        str(selected_path),
    )
    created = create_docx("Рядок один\nРядок два", str(created_path), "Витяг")

    selected_document = Document(saved.path)
    created_document = Document(created.path)
    assert [paragraph.text for paragraph in selected_document.paragraphs] == [
        "Другий пункт для Браво."
    ]
    assert [paragraph.text for paragraph in created_document.paragraphs] == [
        "Витяг",
        "Рядок один",
        "Рядок два",
    ]
    assert selected_document.paragraphs[0].paragraph_format.keep_together is True


def test_registry_exposes_specialized_word_ports():
    registry = NodeRegistry()
    registry.reload()

    reader = registry.get("builtin.word.read_docx")
    saver = registry.get("builtin.word.save_selected_paragraphs")

    assert {port.name: port.data_type for port in reader.outputs} == {
        "document": "WordDocument",
        "file_name": "str",
        "paragraphs": "WordParagraphs",
        "text": "str",
    }
    saver_inputs = {port.name: port.data_type for port in saver.inputs}
    assert saver_inputs["document"] == "WordDocument"
    assert saver_inputs["paragraphs"] == "WordParagraphs"
