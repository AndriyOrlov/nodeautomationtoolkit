import pytest
from nodeautomationtoolkit.builtin_nodes.text_analysis import (
    split_order_blocks,
    group_items_by_markers,
    locate_order_layout_anchors,
    extract_order_fields,
)
from nodeautomationtoolkit.builtin_nodes.recipient_mapping import (
    read_recipient_mapping,
    groups_to_ciphers,
)
from nodeautomationtoolkit.builtin_nodes.word_batch import (
    create_document_batch,
    batch_replace_text,
    batch_keep_items_together,
)
from nodeautomationtoolkit.core.table_types import DataTable


SAMPLE_ORDER_TEXT = """НАКАЗ
командувача військ оперативне командування «Північ»
15 січня 2026 року          м. Чернігів          № 142

По особовому складу

ЗВІЛЬНИТИ І ПРИЗНАЧИТИ до військова частина А1111:
1. Капітана ІВАНОВА Івана Івановича з посади командира роти.
Призначити його на посаду заступника батальйону.

2. Майора ПЕТРОВА Петра Петровича з посади начальника штабу.
Призначити його на посаду командира батальйону.

ПЕРЕМІСТИТИ:
3. Лейтенанта СИДОРОВА Сидора Сидоровича, військова частина А2222, на посаду інструктора у військова частина А1111.

ЗВІЛЬНИТИ:
4. Сержанта КОВАЛЕНКА Василя Васильовича, військова частина А2222, з військової служби у запас.

Командувач військ оперативне командування «Північ»
генерал-майор                                  Олексій ЧЕРНІГІВ

З оригіналом вірно:
Начальник служби діловодства
капітан                                        Сергій ПРАВДІВЕЦЬ
"""


def test_order_layout_anchors_detection():
    anchors_info = locate_order_layout_anchors(SAMPLE_ORDER_TEXT)
    assert anchors_info["date_position"] > 0
    assert anchors_info["number_position"] > 0
    assert isinstance(anchors_info["candidates"], DataTable)


def test_extract_order_fields():
    fields = extract_order_fields(SAMPLE_ORDER_TEXT)
    assert fields["order_number"] == "142"
    assert "15 січня 2026" in fields["order_date"]
    assert len(fields["positions"]) >= 1


def test_split_order_blocks_and_structure():
    split_res = split_order_blocks(SAMPLE_ORDER_TEXT)
    blocks = split_res["blocks"]
    assert len(blocks) >= 4
    assert len(split_res["action_headers"]) >= 3
    assert len(split_res["items"]) >= 4
    assert "командувач" in split_res["signature"].lower()


def test_action_header_inheritance_and_grouping():
    split_res = split_order_blocks(SAMPLE_ORDER_TEXT)
    blocks = split_res["blocks"]

    markers = ["військова частина А1111", "військова частина А2222"]
    group_res = group_items_by_markers(blocks=blocks, markers=markers)
    groups = group_res["groups"]

    assert "військова частина А1111" in groups
    assert "військова частина А2222" in groups

    content_a1111 = groups["військова частина А1111"]
    content_a2222 = groups["військова частина А2222"]

    assert "ІВАНОВА" in content_a1111
    assert "ПЕТРОВА" in content_a1111
    assert "СИДОРОВА" in content_a1111
    assert "СИДОРОВА" in content_a2222
    assert "КОВАЛЕНКА" in content_a2222


def test_recipient_mapping_conversion(tmp_path):
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["Відкрите найменування", "Шифр", "Куди направляється"])
    ws.append(["військова частина А1111", "ВЧ_А1111", "м. Київ"])
    ws.append(["військова частина А2222", "ВЧ_А2222", "м. Львів"])
    excel_path = tmp_path / "mapping.xlsx"
    wb.save(excel_path)

    mapped_res = read_recipient_mapping(str(excel_path))
    assert mapped_res["count"] == 2

    groups = {
        "військова частина А1111": "Текст пункту 1...",
        "військова частина А2222": "Текст пункту 2...",
    }
    converted = groups_to_ciphers(groups=groups, counts={"військова частина А1111": 1, "військова частина А2222": 1}, mapping=mapped_res["mapping"])
    assert "ВЧ_А1111" in converted["documents"]
    assert "ВЧ_А2222" in converted["documents"]
    assert len(converted["documents"]) == 2


def test_word_batch_operations(tmp_path):
    import docx
    doc = docx.Document()
    doc.add_paragraph("ВИТЯГ З НАКАЗУ № {{order_num}}")
    doc.add_paragraph("ЗМІСТ: {{content}}")
    doc_path = tmp_path / "template.docx"
    doc.save(doc_path)

    groups = {
        "ВЧ_А1111": "Пункт 1. Капітан Іванов",
        "ВЧ_А2222": "Пункт 2. Майор Петров",
    }
    batch = create_document_batch(source_path=str(doc_path), groups=groups)
    batch = batch_replace_text(batch, find="{{order_num}}", replacement_text="142")
    batch = batch_keep_items_together(batch)

    assert len(batch.variants) == 2
    assert len(batch.operations) == 2
