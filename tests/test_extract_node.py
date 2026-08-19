import tempfile
from pathlib import Path
import docx

from nodeautomationtoolkit.builtin_nodes.word_batch import create_unit_extracts


def test_create_unit_extracts_structure():
    unit_paragraphs = {
        "160 ОМБр А1600": {
            "unit_code": "160 ОМБр А1600",
            "recipient_to": "Командиру військової частини А1600",
            "destination_where": "м. Львів, вул. Городоцька, 1",
            "header_lines": ["НАКАЗ командира 10 армійського корпусу", "15 серпня 2026 року м. Львів № 77-рс"],
            "items": [
                {
                    "parent_heading": "ПО ОСОБОВОМУ СКЛАДУ:",
                    "label": "1.1.",
                    "text": "1.1. Майора СИДОРЕНКА П.П. направити до військової частини А1600."
                }
            ]
        }
    }

    with tempfile.TemporaryDirectory() as tmp_dir:
        res = create_unit_extracts(
            unit_paragraphs=unit_paragraphs,
            output_folder=tmp_dir,
            order_number="77-рс",
            order_date="15 серпня 2026 року",
            signatory_title="Т.в.о. начальника штабу військової частини А0000",
            signatory_rank_name="майор Петро СИДОРЕНКО",
            certify_extract=True,
            copy_number="прим. 2",
        )

        assert res["count"] == 1
        assert len(res["paths"]) == 1
        assert "combined_path" in res
        assert Path(res["combined_path"]).is_file()

        docx_path = Path(res["paths"][0])
        assert docx_path.is_file()

        doc = docx.Document(docx_path)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "НАКАЗ командира 10 армійського корпусу" in full_text
        assert "Командиру військової частини А1600" in full_text
        assert "м. Львів, вул. Городоцька, 1" in full_text
        assert "ВИТЯГ З НАКАЗУ" in full_text
        assert "ПО ОСОБОВОМУ СКЛАДУ:" in full_text
        assert "Майора СИДОРЕНКА П.П." in full_text
        assert "Згідно з оригіналом:" in full_text
        assert "Т.в.о. начальника штабу військової частини А0000" in full_text
        assert "майор Петро СИДОРЕНКО" in full_text


def test_create_unit_extracts_with_template():
    unit_paragraphs = {
        "160 ОМБр А1600": {
            "unit_code": "160 ОМБр А1600",
            "recipient_to": "Командиру 160 ОМБр",
            "destination_where": "с. Старичі",
            "items": [
                {
                    "parent_heading": "НАПРАВИТИ:",
                    "text": "1. Лейтенанта Коваля направити до в/ч А1600."
                }
            ]
        }
    }

    with tempfile.TemporaryDirectory() as tmp_dir:
        template_file = Path(tmp_dir) / "template.docx"
        tmpl_doc = docx.Document()
        tmpl_doc.add_paragraph("КОМУ: {{кому}}")
        tmpl_doc.add_paragraph("КУДИ: {{куди}}")
        tmpl_doc.add_paragraph("ВИТЯГ З НАКАЗУ № {{номер_наказу}} ВІД {{дата_наказу}}")
        tmpl_doc.add_paragraph("{{зміст}}")
        tmpl_doc.add_paragraph("{{засвідчення}}")
        tmpl_doc.save(template_file)

        res = create_unit_extracts(
            unit_paragraphs=unit_paragraphs,
            output_folder=tmp_dir,
            template_docx_path=str(template_file),
            order_number="105",
            order_date="20 серпня 2026 року",
            signatory_title="Начальник штабу в/ч А0000",
            signatory_rank_name="полковник Іван ІВАНОВ",
            certify_extract=True,
            copy_number="прим. 1",
        )

        assert res["count"] == 1
        docx_path = Path(res["paths"][0])
        doc = docx.Document(docx_path)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "КОМУ: Командиру 160 ОМБр" in full_text
        assert "КУДИ: с. Старичі" in full_text
        assert "ВИТЯГ З НАКАЗУ № 105 ВІД 20 серпня 2026 року" in full_text
        assert "1. Лейтенанта Коваля направити до в/ч А1600." in full_text
        assert "Згідно з оригіналом:" in full_text
        assert "Начальник штабу в/ч А0000" in full_text
        assert "полковник Іван ІВАНОВ" in full_text
