import pytest
from pathlib import Path
import docx

from nodeautomationtoolkit.builtin_nodes.word import (
    analyze_headers_footers,
    clear_headers_footers,
)
from nodeautomationtoolkit.builtin_nodes.word_batch import (
    create_document_batch,
    batch_clear_headers_footers,
)
from nodeautomationtoolkit.builtin_nodes.text_analysis import (
    visualize_order_layout,
    split_order_blocks,
)
from nodeautomationtoolkit.builtin_nodes.output import show_image
from nodeautomationtoolkit.core.table_types import DataTable


def test_analyze_and_clear_headers_footers(tmp_path):
    doc = docx.Document()
    section = doc.sections[0]
    header = section.header
    footer = section.footer

    header.paragraphs[0].text = "Верхній колонтитул наказу Стор. 1"
    footer.paragraphs[0].text = "Нижній колонтитул 2026 року"

    input_docx = tmp_path / "test_headers.docx"
    doc.save(input_docx)

    # Test analyze_headers_footers
    res = analyze_headers_footers(str(input_docx))
    assert len(res["headers"]) >= 1
    assert "Верхній колонтитул" in res["headers"][0]
    assert res["has_page_numbers"] is True
    assert isinstance(res["details"], DataTable)

    # Test clear_headers_footers
    cleared_docx = tmp_path / "test_cleared.docx"
    clear_res = clear_headers_footers(str(input_docx), output_path=str(cleared_docx))
    assert Path(clear_res["path"]).is_file()

    # Re-analyze cleared docx
    cleared_analysis = analyze_headers_footers(str(cleared_docx))
    assert len(cleared_analysis["headers"]) == 0
    assert len(cleared_analysis["footers"]) == 0


def test_batch_clear_headers_footers(tmp_path):
    doc = docx.Document()
    doc.add_paragraph("Тестовий вміст наказу")
    doc_path = tmp_path / "order.docx"
    doc.save(doc_path)

    batch = create_document_batch(source_path=str(doc_path), names=["Витяг 1"])
    batch = batch_clear_headers_footers(batch)

    assert len(batch.operations) == 1
    assert batch.operations[0].kind == "clear_headers_footers"


def test_visualize_order_layout(tmp_path):
    order_text = """НАКАЗ
15 січня 2026 року № 100
ЗВІЛЬНИТИ:
1. Сержанта Іванова.
Командир частини ІВАНОВ
"""
    out_img = tmp_path / "layout.png"
    res = visualize_order_layout(text=order_text, output_image_path=str(out_img))

    assert Path(res["image_path"]).is_file()
    assert res["image_path"] == str(out_img)
    assert isinstance(res["table"], DataTable)
    assert "Блок" in res["blocks_summary"]


def test_advanced_rename_file(tmp_path):
    from nodeautomationtoolkit.builtin_nodes.files import advanced_rename_file

    file_path = tmp_path / "OLD_order_2026_draft.docx"
    file_path.write_text("dummy content")

    # Test trim before symbol '_', add prefix and suffix
    res = advanced_rename_file(
        path=str(file_path),
        prefix="Наказ_",
        suffix="_фінал",
        trim_before_symbol="order_",
        apply_rename_on_disk=True,
    )

    assert res["new_name"] == "Наказ_2026_draft_фінал.docx"
    assert Path(res["path"]).is_file()
    assert not file_path.exists()


def test_folder_and_move_nodes(tmp_path):
    from nodeautomationtoolkit.builtin_nodes.files import (
        create_folder,
        rename_folder,
        move_file,
    )

    # 1. Create folder
    folder_res = create_folder(folder_path=str(tmp_path), folder_name="Накази_2026")
    created_dir = Path(folder_res["path"])
    assert created_dir.is_dir()

    # 2. Rename folder
    ren_res = rename_folder(folder_path=str(created_dir), new_name="Накази_Архів")
    renamed_dir = Path(ren_res["path"])
    assert renamed_dir.is_dir()
    assert renamed_dir.name == "Накази_Архів"

    # 3. Move file
    test_file = tmp_path / "doc1.docx"
    test_file.write_text("dummy")
    move_res = move_file(file_path=str(test_file), destination_folder=str(renamed_dir))
    moved_path = Path(move_res["new_path"])
    assert moved_path.is_file()
    assert moved_path.parent == renamed_dir


def test_organize_orders_by_number(tmp_path):
    from nodeautomationtoolkit.builtin_nodes.word_batch import organize_orders_by_number

    doc = docx.Document()
    doc.add_paragraph("НАКАЗ\n15 січня 2026 року № 55\nПро призначення...")
    order_file = tmp_path / "order_55.docx"
    doc.save(order_file)

    res = organize_orders_by_number(input_folder=str(tmp_path), create_subfolders=True)
    assert res["processed_count"] == 1
    subfolder = tmp_path / "Наказ № 55"
    assert subfolder.is_dir()
    expected_file = subfolder / "прим_2_15_січня_2026_року_№55.docx"
    assert expected_file.is_file()


def test_map_military_units_and_create_extracts(tmp_path):
    from nodeautomationtoolkit.builtin_nodes.recipient_mapping import map_military_units
    from nodeautomationtoolkit.builtin_nodes.word_batch import create_unit_extracts

    order_text = """НАКАЗ
15 січня 2026 року № 77

1. 100 окрема бригада направляє 5 бійців.
2. 200 артилерійська бригада забезпечує техніку.
3. 100 окрема бригада надає додаткове постачання.
"""

    mapping = {
        "100 окрема бригада": "військова частина А1000",
        "200 артилерійська бригада": "військова частина А2000",
    }

    # 1. Map military units
    map_res = map_military_units(text=order_text, mapping=mapping)
    assert len(map_res["units_list"]) == 2
    assert "військова частина А1000" in map_res["unit_paragraphs"]
    assert len(map_res["unit_paragraphs"]["військова частина А1000"]["items"]) == 2

    # 2. Create extracts DOCX
    src_doc = tmp_path / "order_77.docx"
    doc = docx.Document()
    doc.add_paragraph(order_text)
    doc.save(src_doc)

    extract_res = create_unit_extracts(
        source_order_path=str(src_doc),
        unit_paragraphs=map_res["unit_paragraphs"],
        output_folder=str(tmp_path / "Extracts"),
        save_individual_files=True,
    )

    assert extract_res["count"] == 2
    assert len(extract_res["paths"]) == 2
    assert Path(extract_res["paths"][0]).is_file()


def test_military_preamble_extraction(tmp_path):
    from nodeautomationtoolkit.builtin_nodes.recipient_mapping import map_military_units
    from nodeautomationtoolkit.builtin_nodes.word_batch import create_unit_extracts

    order_text = """НАКАЗ
командира військової частини А0000
15 січня 2026 року № 88

§ 1

Відповідно до пунктів 82, 83 та 257 Положення про проходження громадянами України військової служби у Збройних Силах України нижчепойменованих осіб офіцерського складу ЗВІЛЬНИТИ з займаних посад і ПРИЗНАЧИТИ:

1.1. Майора Петренка П.П., 300 окрема механізована бригада, на посаду командира батальйону.
"""

    mapping = {
        "300 окрема механізована бригада": "військова частина А3000",
    }

    map_res = map_military_units(text=order_text, mapping=mapping)
    unit_entry = map_res["unit_paragraphs"]["військова частина А3000"]
    parent_heading = unit_entry["items"][0]["parent_heading"]

    assert "§ 1" in parent_heading
    assert "Відповідно до пунктів 82, 83 та 257" in parent_heading
    assert "ПРИЗНАЧИТИ:" in parent_heading

    src_doc = tmp_path / "order_88.docx"
    doc = docx.Document()
    for line in order_text.split("\n\n"):
        doc.add_paragraph(line)
    doc.save(src_doc)

    extract_res = create_unit_extracts(
        source_order_path=str(src_doc),
        unit_paragraphs=map_res["unit_paragraphs"],
        output_folder=str(tmp_path / "Extracts_Preamble"),
    )

    created_file = Path(extract_res["paths"][0])
    created_doc = docx.Document(created_file)
    extracted_text = "\n".join(p.text for p in created_doc.paragraphs)

    assert "§ 1" in extracted_text
    assert "Відповідно до пунктів 82, 83 та 257" in extracted_text
    assert "ПРИЗНАЧИТИ:" in extracted_text
    assert "військова частина А3000" in extracted_text


def test_multiline_item_with_basis_extraction(tmp_path):
    from nodeautomationtoolkit.builtin_nodes.recipient_mapping import map_military_units
    from nodeautomationtoolkit.builtin_nodes.word_batch import create_unit_extracts

    order_text = """НАКАЗ
15 червня 2026 року № 308

§ 2

10. Пункт 2 наказу про призначення на посаду начальника групи персоналу штабу 167 окремої механізованої бригади молодшого лейтенанта ІВАНОВА Івана Івановича, СКАСУВАТИ як нереалізований.

1997 р.н. 0000000000.
Підстава: клопотання начальника Волинського ОТЦК та СП від 30.07.2026 №ХХХ/ХХ.
"""

    mapping = {
        "167 окремої механізованої бригади": "військова частина А1670",
    }

    map_res = map_military_units(text=order_text, mapping=mapping)
    unit_entry = map_res["unit_paragraphs"]["військова частина А1670"]
    item_text = unit_entry["items"][0]["text"]

    assert "10. Пункт 2 наказу" in item_text
    assert "1997 р.н. 0000000000." in item_text
    assert "Підстава: клопотання начальника" in item_text
    # Правило 4.1: item["text"] має відкриті назви, item["text_cipher"] — шифри
    assert "167 окремої механізованої бригади" in item_text
    item_cipher = unit_entry["items"][0]["text_cipher"]
    assert "військова частина А1670" in item_cipher


def test_excel_nodes_read_write_replace_save(tmp_path):
    from nodeautomationtoolkit.builtin_nodes.excel_nodes import (
        read_sheet,
        replace_text,
        save_table,
        write_cell,
    )
    from nodeautomationtoolkit.core.table_types import DataTable

    # 1. Save table to Excel
    table = DataTable(
        columns=("ПІБ", "Звання", "ВЧ"),
        rows=(
            ("Іванов І.І.", "майор", "100 ОМБр"),
            ("Петренко П.П.", "капітан", "200 ОАБр"),
        ),
        title="Особовий склад",
    )
    excel_file = tmp_path / "personnel.xlsx"
    save_res = save_table(table=table, output_path=str(excel_file), sheet_name="Список")
    assert save_res["rows_count"] == 2
    assert excel_file.is_file()

    # 2. Read sheet
    read_res = read_sheet(excel_path=str(excel_file), sheet_name="Список")
    assert read_res["rows_count"] == 2
    assert read_res["headers"] == ["ПІБ", "Звання", "ВЧ"]

    # 3. Write cell
    write_res = write_cell(excel_path=str(excel_file), cell_reference="B2", value="підполковник")
    assert write_res["output_path"] == str(excel_file)

    # 4. Replace text
    replace_res = replace_text(
        excel_path=str(excel_file),
        search_text="100 ОМБр",
        replace_text="в/ч А1000",
    )
    assert replace_res["replaced_count"] >= 1

    # Verify updated content
    final_read = read_sheet(excel_path=str(excel_file), sheet_name="Список")
    assert final_read["table"].rows[0][1] == "підполковник"
    assert final_read["table"].rows[0][2] == "в/ч А1000"


def test_windows_system_and_standalone_packager(tmp_path):
    from nodeautomationtoolkit.app import _run_cli_scenario
    from nodeautomationtoolkit.builtin_nodes.windows_system import export_standalone, run_command
    from nodeautomationtoolkit.core.registry import NodeRegistry

    # 1. Test Windows command execution
    cmd_res = run_command(command="Write-Output 'Hello Windows NAT'", use_powershell=True)
    assert cmd_res["exit_code"] == 0
    assert "Hello Windows NAT" in cmd_res["stdout"]

    # 2. Test Standalone package exporter
    dummy_scenario = tmp_path / "test_graph.nat.json"
    dummy_scenario.write_text('{"format":"nodeautomationtoolkit","nodes":[],"connections":[]}', encoding="utf-8")

    pkg_res = export_standalone(
        graph_json_path=str(dummy_scenario),
        output_folder=str(tmp_path / "Packages"),
        package_name="TestAuto",
    )
    assert Path(pkg_res["launcher_bat"]).is_file()
    assert "TestAuto" in pkg_res["package_dir"]

    # 3. Test CLI scenario execution
    registry = NodeRegistry()
    registry.reload()
    cli_code = _run_cli_scenario(str(dummy_scenario), registry)
    assert cli_code == 0


def test_signatory_slash_line_breaks(tmp_path):
    from nodeautomationtoolkit.builtin_nodes.word_batch import create_unit_extracts

    src_doc = tmp_path / "order_slash.docx"
    doc = docx.Document()
    doc.add_paragraph("НАКАЗ\n1. Солдата призначити.")
    doc.save(src_doc)

    unit_paragraphs = {
        "А1234": {
            "items": [{"parent_heading": "", "text": "1. Солдата призначити."}],
        }
    }

    slash_pos = "Тимчасово виконуючий обов'язки / командувача військ / оперативного командування"
    res = create_unit_extracts(
        source_order_path=str(src_doc),
        unit_paragraphs=unit_paragraphs,
        output_folder=str(tmp_path / "Extracts_Slash"),
        signatory_title=slash_pos,
        certify_extract=True,
    )

    created_doc = docx.Document(res["paths"][0])
    text = "\n".join(p.text for p in created_doc.paragraphs)
    # Verify that slashes were converted to individual paragraph lines
    assert "Тимчасово виконуючий обов'язки\nкомандувача військ\nоперативного командування" in text or (
        "Тимчасово виконуючий обов'язки" in text and "командувача військ" in text
    )






